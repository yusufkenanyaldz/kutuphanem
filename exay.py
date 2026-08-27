"""
e-YMM Karşıt İnceleme — KDV Fatura Listesi Bölme Programı v5.1
Kapsam kriterleri:
  1) Tek fatura >= esik_tek  VEYA  toplam >= esik_toplam
  2) Seçilenler listenin %80'ini karşılamıyorsa kalan firmalar
     büyükten küçüğe eklenir (tüm faturalarıyla birlikte)
Seri numarası: tarihten sonraki ayrı sütundan alınır, yoksa boş
Fatura numarası: olduğu gibi (harfler dahil)
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading, os, re, sys, json
from datetime import datetime
from pathlib import Path

# Kullanıcının son kullandığı eşik/yüzde değerleri burada saklanır ki program
# her açılışta aynı kriterlerle gelsin. Program Files gibi yazılamayan yerlere
# kurulmuş .exe'de sorun olmaması için kullanıcı ana klasörüne yazılır.
AYAR_YOLU = Path.home() / '.exay_ayarlar.json'


def kaynak_yolu(rel_yol):
    """Hem normal çalışmada hem de PyInstaller ile paketlenmiş .exe içinde
    (logo gibi) yardımcı dosyaların doğru yolunu döndürür."""
    if getattr(sys, 'frozen', False):          # .exe olarak paketlenmişse
        taban = Path(getattr(sys, '_MEIPASS', Path(sys.executable).parent))
    else:                                       # normal .py çalışması
        taban = Path(__file__).parent
    return taban / rel_yol

try:
    import pandas as pd
    import openpyxl
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable,'-m','pip','install',
                           'pandas','openpyxl','xlrd','--quiet'])
    import pandas as pd, openpyxl

SURUM = "5.1"   # GUI başlığında ve özet raporda gösterilir

SABLON_SUTUNLAR = [
    "Faturanın Tarihi","Faturanın Serisi","Faturanın Numarası",
    "Faturanın Tutarı (TL)","K.D.V(TL)","Defter Kayıt Tarihi",
    "Yevmiye Numarası","Ödeme Şekli ve Ödemeye İlişkin Bilgiler",
    "Açıklama","Hatalı Satır Açıklama",
]

# ══════════════════════════════════════════
#  YARDIMCI
# ══════════════════════════════════════════
def sutun_bul(kolonlar, aranacak):
    for col in kolonlar:
        cs = str(col).lower().strip()
        for a in aranacak:
            if a.lower() in cs: return col
    return None

def kdv_sutunu_bul(kolonlar):
    """Faturanın KDV'si sütununu bulur. 'KDV'si', 'KDV si', 'KDVsi', 'KDV' gibi
    biçimlerin hepsini yakalar; 'KDV Hariç Tutarı' (matrah), 'Toplam İndirilecek KDV',
    tevkifat / 2 nolu beyanname KDV sütunlarıyla KARIŞMAZ.
    Böylece hem eski (KDV'si) hem yeni (KDV si) liste tipinde doğru sütun seçilir."""
    yasak = ('hariç', 'haric', 'toplam', 'tevkifat', '2 nolu', 'ödenen', 'odenen',
             'indirilecek', 'indirilen', 'dönem', 'donem', 'matrah', 'tutarı', 'tutari')
    # 1) Dar eşleşme: sadeleştirince 'kdvsi' veya 'kdv' olan sütun (asıl KDV'si sütunu)
    for col in kolonlar:
        cs = str(col).lower().strip()
        if 'kdv' not in cs or any(y in cs for y in yasak):
            continue
        temiz = cs.replace("'", '').replace(' ', '').replace('’', '')
        if temiz in ('kdvsi', 'kdv'):
            return col
    # 2) Gevşek son çare: 'kdv' içeren ama yasaklı kelime içermeyen ilk sütun
    for col in kolonlar:
        cs = str(col).lower().strip()
        if 'kdv' in cs and not any(y in cs for y in yasak):
            return col
    return None

def seri_sutunu_bul(kolonlar, tarih_col, faturano_col=None):
    """
    Fatura serisi sütununu bul:
      1) Adında 'seri' geçen bir sütun varsa doğrudan onu kullan.
      2) Yoksa tarih sütununun hemen sağındaki sütunu aday al — ancak bu sütun
         fatura NUMARASI sütununun ta kendisiyse seri sütunu yok demektir; boş bırak.
    Böylece tarihin sağında ayrı bir seri sütunu olmayan listelerde
    fatura numarası yanlışlıkla 'seri' alanına yazılmaz.
    """
    # 1) Adında "seri" geçen sütun (en güvenilir)
    for col in kolonlar:
        if 'seri' in str(col).lower():
            return col
    # 2) Tarihin sağındaki sütun — ama numara sütunu değilse
    if tarih_col is None: return None
    try:
        idx = list(kolonlar).index(tarih_col)
        if idx + 1 < len(kolonlar):
            aday = kolonlar[idx + 1]
            if faturano_col is not None and aday == faturano_col:
                return None   # sağdaki sütun aslında numara → seri yok
            return aday
    except: pass
    return None

def _csv_okuyucu_hazirla(dosya):
    """CSV/TXT dosyaları için kodlama ve sütun ayracını otomatik saptayıp,
    read_excel ile aynı arayüzde (header/skiprows kabul eden) bir okuyucu
    döndürür. Türkçe muhasebe çıktıları çoğunlukla ';' ayraçlı ve Windows
    (cp1254) kodludur; UTF-8 ve virgül/tab de denenir. dtype=str ile VKN ve
    fatura no'daki baştaki sıfırlar korunur (para_deger tutarları yine ayrıştırır)."""
    enc = 'utf-8'
    ornek = ''
    for e in ('utf-8-sig', 'utf-8', 'cp1254', 'iso-8859-9'):
        try:
            with open(dosya, 'r', encoding=e) as f:
                ornek = f.read(8192)
            enc = e
            break
        except (UnicodeDecodeError, LookupError):
            continue
    sayac = {s: ornek.count(s) for s in (';', '\t', ',')}
    sep = max(sayac, key=sayac.get) if any(sayac.values()) else ';'
    return lambda **kw: pd.read_csv(dosya, sep=sep, encoding=enc,
                                    dtype=str, **kw)

def ana_listeyi_oku(dosya):
    ext = Path(dosya).suffix.lower()
    if ext in ('.csv', '.txt'):
        okuyucu = _csv_okuyucu_hazirla(dosya)
    else:
        engine = 'xlrd' if ext == '.xls' else 'openpyxl'
        okuyucu = lambda **kw: pd.read_excel(dosya, engine=engine, **kw)
    raw = okuyucu(header=None)

    # Başlık satırını bul — birden fazla anahtar kelimeyle dene
    ANAHTAR = ['alış faturası', 'satıcının', 'vergi kimlik',
               'fatura tarihi', 'faturanın tarihi', 'kdv hariç']
    baslik = None
    for i, row in raw.iterrows():
        v = ' '.join(str(x) for x in row if pd.notna(x)).lower()
        if any(a in v for a in ANAHTAR):
            baslik = i; break

    if baslik is None:
        # Son çare: en fazla dolu hücre içeren satırı başlık say
        dolu = raw.apply(lambda r: r.notna().sum(), axis=1)
        baslik = int(dolu.idxmax())

    df = okuyucu(skiprows=baslik, header=0)
    df = df.dropna(how='all').reset_index(drop=True)

    # Başlık satırı tekrar veri olarak geldiyse at
    if df.iloc[0].astype(str).str.contains(
            'Alış Faturası|Satıcın|Vergi Kimlik|Fatura Tarihi', na=False).any():
        df = df.iloc[1:].reset_index(drop=True)

    # Sütun adlarını normalize et (baştaki/sondaki boşluk, satır sonu)
    df.columns = [str(c).strip().replace('\n', ' ') for c in df.columns]

    # Muhasebe (191 hesabı) dökümü tipiyse sütunları standart GİB adlarına çevir
    df = _muhasebe_tipini_esle(df)
    return df


def _muhasebe_tipini_esle(df):
    """Muhasebe programından alınan 191 hesabı dökümünü tanır ve sütunlarını
    standart GİB listesi adlarına çevirir; böylece programın geri kalanı
    hiçbir değişiklik gerekmeden çalışır.
    Tanıma: 'Borç' + 'vergi kimlik' sütunları var, 'Satıcının...' sütunu yok.
    Eşleme: Tarih→Alış Faturasının Tarihi, fatura no→Sıra No'su,
            Açıklama→Satıcı Ünvanı, Borç→KDV'si,
            matrah / KDV'nin sağındaki adsız sayısal sütun→KDV Hariç Tutar."""
    kolonlar = list(df.columns)
    kl = [str(c).lower().strip() for c in kolonlar]
    var = lambda k: any(k in c for c in kl)
    if not (var('borç') or var('borc')):   return df   # muhasebe tipi değil
    if not var('vergi kimlik'):            return df
    if var('satıcının') or var('saticinin'): return df # zaten GİB tipi

    esle = {}
    for orij, cl in zip(kolonlar, kl):
        if   cl == 'tarih':                 esle[orij] = 'Alış Faturasının Tarihi'
        elif cl.startswith('fatura no'):    esle[orij] = "Alış Faturasının Sıra No'su"
        elif cl in ('açıklama', 'aciklama'):esle[orij] = 'Satıcının Adı-Soyadı / Ünvanı'
        elif cl in ('borç', 'borc'):        esle[orij] = "KDV'si"
        elif 'matrah' in cl:                esle[orij] = 'Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı'
    df = df.rename(columns=esle)

    # Matrah sütunu adsızsa: KDV'nin hemen sağındaki, çoğunluğu sayısal sütun matrahtır
    MATRAH = 'Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı'
    if MATRAH not in df.columns and "KDV'si" in df.columns:
        cols = list(df.columns)
        i = cols.index("KDV'si")
        if i + 1 < len(cols):
            aday = cols[i + 1]
            vals = pd.to_numeric(df[aday].apply(para_deger), errors='coerce')
            if len(vals) and vals.notna().mean() > 0.6:
                df = df.rename(columns={aday: MATRAH})
    return df

def tarih_fmt(val):
    if pd.isna(val) or str(val).strip() == '': return ''
    if isinstance(val, datetime): return val.strftime('%d.%m.%Y')
    s = str(val).strip()
    try: return datetime.strptime(s[:10], '%Y-%m-%d').strftime('%d.%m.%Y')
    except: return s

def sayi_fmt(val):
    if pd.isna(val) or str(val).strip() == '': return ''
    try:
        f = float(str(val).replace(',', '.'))
        return str(int(f)) if f == int(f) else f'{f:.2f}'.replace('.', ',')
    except: return str(val).strip()

def para_deger(val):
    """Kaynaktaki tutarı gerçek sayıya (float) çevirir; küsürat korunur.
    float, '45927,50', '1.234.567,89', '1,234,567.89' gibi girişleri destekler.
    Boş/geçersizse None döner."""
    if pd.isna(val) or str(val).strip() == '':
        return None
    if isinstance(val, (int, float)):
        return round(float(val), 2)
    s = str(val).strip().replace(' ', '').replace('₺', '').replace('TL', '')
    if ',' in s and '.' in s:
        # Hem nokta hem virgül varsa: en sağdaki ondalık ayracıdır
        if s.rfind(',') > s.rfind('.'):
            s = s.replace('.', '').replace(',', '.')   # 1.234,56 -> 1234.56
        else:
            s = s.replace(',', '')                     # 1,234.56 -> 1234.56
    elif ',' in s:
        s = s.replace(',', '.')                        # 1234,56 -> 1234.56
    try:
        return round(float(s), 2)
    except:
        return None

def para_oku(val):
    try: return float(str(val).replace(',', '.').replace(' ', ''))
    except: return 0.0

def dosya_adi_temizle(m): return re.sub(r'[\\/*?:"<>|]', '_', str(m)).strip()[:45]

def donem_bul(stem, df=None):
    """Dönemi (AA.YYYY) tespit eder, öncelik sırasıyla:
       1) Dosya adındaki Türkçe ay adı (+ varsa yıl)
       2) Dosya adındaki sayısal ay + yıl (ör: 04_2026)
       3) Listedeki fatura tarihlerinin en sık görülen ay/yılı
       4) Bugünün ay/yılı (son çare)"""
    # Türkçe karakterleri ASCII'ye indir: kullanıcılar dosya adında ay adını
    # çoğu zaman diakritiksiz yazar (NISAN, SUBAT, AGUSTOS, EYLUL, ARALIK...).
    # Böyle yazımların da doğru aya eşlenmesi için karşılaştırmayı sadeleştir.
    def _ascii_fold(s):
        tr = {'İ':'I','I':'I','Ş':'S','Ğ':'G','Ü':'U','Ö':'O','Ç':'C'}
        return ''.join(tr.get(c, c) for c in s.upper())
    fn = _ascii_fold(stem)
    aylar = {'OCAK':'01','ŞUBAT':'02','MART':'03','NİSAN':'04',
              'MAYIS':'05','HAZİRAN':'06','TEMMUZ':'07','AĞUSTOS':'08',
              'EYLÜL':'09','EKİM':'10','KASIM':'11','ARALIK':'12'}
    yil_m = re.search(r'(20\d{2})', fn)
    for ad, no in aylar.items():
        if _ascii_fold(ad) in fn:
            return f"{no}.{yil_m.group(1) if yil_m else datetime.now().year}"
    # Sayısal ay (tek başına 01-12; uzun sayı dizilerinin içi sayılmaz)
    ay_m = re.search(r'(?<!\d)(0[1-9]|1[0-2])(?!\d)', fn)
    if ay_m and yil_m:
        return f"{ay_m.group(1)}.{yil_m.group(1)}"
    # Veri tarihlerinden: en sık görülen ay/yıl
    if df is not None:
        tarih_col = sutun_bul(list(df.columns),
                              ['alış faturasının tarihi', 'fatura tarihi', 'tarih'])
        if tarih_col is not None:
            try:
                t = pd.to_datetime(df[tarih_col], errors='coerce', dayfirst=True).dropna()
                if len(t):
                    mod = t.dt.strftime('%m.%Y').mode()
                    if len(mod):
                        return mod.iloc[0]
            except Exception:
                pass
    if ay_m:
        return f"{ay_m.group(1)}.{datetime.now().year}"
    return datetime.now().strftime('%m.%Y')

def _gecersizlik_nedeni(x) -> str:
    s = str(x).strip().replace('.0', '').replace(' ', '')
    if s == '' or s.lower() in ('nan', 'none', 'nat'):
        return 'Boş/tanımsız'
    if not s.isdigit():
        return f'Sayısal değil: "{s[:20]}"'
    if len(s) < 10:
        return f'Çok kısa ({len(s)} hane, min 10)'
    if len(s) > 11:
        return f'Çok uzun ({len(s)} hane, max 11)'
    if s == s[0] * len(s):
        return f'Yer tutucu/sahte kimlik ({s})'
    return 'Geçersiz'

# ══════════════════════════════════════════
#  KRİTER DOĞRULAMA & DOĞRULUK KONTROLLERİ
#  (hepsi saf/yan etkisiz — yalnızca veri döndürür, iş kuralını değiştirmez)
# ══════════════════════════════════════════
def kriter_dogrula(esik_tek, esik_toplam, yuzde80):
    """Kullanıcının girdiği kriterleri mantıklı aralıkta mı diye denetler.
    (ok: bool, mesaj: str) döndürür; ok False ise mesaj kullanıcıya gösterilecek
    Türkçe hatadır. İş kuralını değiştirmez, yalnızca hatalı girişi engeller."""
    try:
        et, eto, y = float(esik_tek), float(esik_toplam), float(yuzde80)
    except (TypeError, ValueError):
        return (False, "Kriter değerleri sayı olmalıdır.")
    if et <= 0 or eto <= 0:
        return (False, "Fatura limitleri 0'dan büyük olmalıdır.")
    if not (0 < y <= 100):
        return (False, "Kapsam yüzdesi 0 ile 100 arasında olmalıdır.")
    return (True, "")

def bulunan_sutunlar(df):
    """İşlem öncesi önizleme için: kritik alanların hangi başlıklara eşlendiğini
    (etiket → sütun adı / None) sözlüğü olarak döndürür."""
    K = list(df.columns)
    return {
        'VKN/TC':    sutun_bul(K, ['vergi kimlik', 'vkn', 'tc kimlik']),
        'Tarih':     sutun_bul(K, ['alış faturasının tarihi', 'fatura tarihi', 'tarih']),
        'Fatura No': sutun_bul(K, ['alış faturasının sıra no', 'fatura no']),
        'Matrah':    sutun_bul(K, ['kdv hariç tutarı', 'faturanın tutarı']),
        'KDV':       kdv_sutunu_bul(K) or sutun_bul(K, ['toplam indirilecek kdv']),
        'Ünvan':     sutun_bul(K, ['satıcının adı', 'ünvanı', 'unvan']),
    }

def kdv_tutarlilik_kontrol(df):
    """KDV sütunu ile matrah sütununun oranına bakar; oran makul KDV oranlarından
    (yaklaşık %1–%20) belirgin uzaksa, muhtemel yanlış sütun eşleşmesine karşı
    uyarı üretir. [(mesaj, tip), ...] döndürür; sorun yoksa boş liste.
    Yalnızca uyarır — seçimi/iş kuralını etkilemez."""
    K = list(df.columns)
    matrah_col = sutun_bul(K, ['kdv hariç tutarı', 'faturanın tutarı'])
    kdv_col = kdv_sutunu_bul(K) or sutun_bul(K, ['toplam indirilecek kdv'])
    if not matrah_col or not kdv_col or matrah_col == kdv_col:
        return []
    oranlar = []
    for mv, kv in zip(df[matrah_col].apply(para_deger), df[kdv_col].apply(para_deger)):
        if mv and kv is not None and mv > 0:
            oranlar.append(kv / mv)
    if not oranlar:
        return []
    oranlar.sort()
    medyan = oranlar[len(oranlar) // 2]
    makul = (0.01, 0.08, 0.10, 0.18, 0.20)
    if min(abs(medyan - r) for r in makul) > 0.03:
        return [(f"  ⚠️  KDV/matrah oranı medyanı %{medyan*100:.1f} — beklenen KDV "
                 f"oranlarına (%1–%20) uzak. KDV ve matrah sütunlarının doğru "
                 f"eşleştiğini kontrol edin.", "warn")]
    return []

def mukerrer_fatura_bul(df):
    """Aynı (VKN, fatura no) ikilisinin birden çok satırda geçtiği mükerrer
    kayıtları bulur. [(vkn, fatura_no, adet), ...] döndürür (adet > 1).
    Yalnızca uyarı amaçlıdır; satırları silmez."""
    K = list(df.columns)
    vkn_col = sutun_bul(K, ['vergi kimlik', 'vkn', 'tc kimlik'])
    fno_col = sutun_bul(K, ['alış faturasının sıra no', 'fatura no'])
    if not vkn_col or not fno_col:
        return []
    say = {}
    for v, f in zip(df[vkn_col], df[fno_col]):
        vs = str(v).strip().replace('.0', '')
        fs = str(f).strip()
        if vs.lower() in ('', 'nan', 'none', 'nat') or fs.lower() in ('', 'nan', 'none', 'nat'):
            continue
        say[(vs, fs)] = say.get((vs, fs), 0) + 1
    return [(v, f, a) for (v, f), a in say.items() if a > 1]

def _ay_yil(val):
    """Bir tarih değerinden 'AA.YYYY' döndürür; ayrıştırılamazsa None.
    tarih_fmt ile aynı biçim önceliğini (önce ISO, sonra gün.ay.yıl) izler ki
    bu kod tabanının tarih düzeniyle tutarlı ve tek anlamlı olsun."""
    if pd.isna(val):
        return None
    if isinstance(val, datetime):      # pd.Timestamp da datetime alt-sınıfıdır
        return val.strftime('%m.%Y')
    s = str(val).strip()
    if not s:
        return None
    for fmt in ('%Y-%m-%d', '%d.%m.%Y', '%d/%m/%Y', '%Y.%m.%d'):
        try:
            return datetime.strptime(s[:10], fmt).strftime('%m.%Y')
        except ValueError:
            continue
    try:                               # son çare: pandas (TR için gün-önce)
        d = pd.to_datetime(s, errors='coerce', dayfirst=True)
        return None if pd.isna(d) else d.strftime('%m.%Y')
    except Exception:
        return None

def donem_disi_tarih_kontrol(df, donem):
    """Fatura tarihlerinden kaçının seçilen dönem (AA.YYYY) ay/yılı DIŞINDA
    olduğunu (toplam_tarihli, donem_disi) olarak döndürür. Yanlış dönem dosyasını
    yakalamaya yarayan bir uyarı ölçütüdür. Tarih sütunu yoksa (0, 0)."""
    K = list(df.columns)
    tarih_col = sutun_bul(K, ['alış faturasının tarihi', 'fatura tarihi', 'tarih'])
    if not tarih_col:
        return (0, 0)
    aylar = [a for a in (_ay_yil(v) for v in df[tarih_col]) if a]
    if not aylar:
        return (0, 0)
    disi = sum(1 for a in aylar if a != donem)
    return (len(aylar), disi)

# ══════════════════════════════════════════
#  FİLTRELEME
# ══════════════════════════════════════════
def firmalari_filtrele(df, esik_tek, esik_toplam, yuzde80, log_cb):
    vkn_col   = sutun_bul(list(df.columns), ['vergi kimlik', 'vkn', 'tc kimlik'])
    tutar_col = sutun_bul(list(df.columns), ['kdv hariç tutarı', 'faturanın tutarı'])
    unvan_col = sutun_bul(list(df.columns), ['satıcının adı', 'ünvanı', 'unvan'])

    if not vkn_col:
        raise ValueError("VKN sütunu bulunamadı.")

    # Geçerli / geçersiz satır ayrımı
    def normalize_vkn(x):
        """VKN/TCKN normalize et: önde sıfır eksikse tamamla"""
        s = str(x).strip().replace('.0', '').replace(' ', '')
        if not s.isdigit(): return s
        # 8-9 hane → 10 haneye (VKN önde sıfır eksik)
        if len(s) in (8, 9):
            s = s.zfill(10)
        return s

    def gecerli_kimlik(x):
        s = normalize_vkn(x)
        if not (s.isdigit() and len(s) in (10, 11)):
            return False
        # Yer tutucu / sahte kimlikler: 0000000000, 1111111111, 2222222222 ...
        # (muhasebe dökümlerinde VKN'si girilmemiş satırlar için kullanılır;
        #  ayrı satıcıları tek sahte firmada birleştirmemek için geçersiz sayılır)
        if s == s[0] * len(s):
            return False
        return True

    # Önce normalize et, sonra geçerli/geçersiz ayır
    df = df.copy()
    df[vkn_col] = df[vkn_col].apply(normalize_vkn)

    mask_gecerli  = df[vkn_col].apply(gecerli_kimlik)
    df_t          = df[mask_gecerli].copy()
    df_gecersiz   = df[~mask_gecerli].copy()   # atlanacak satırlar

    df_t[vkn_col] = df_t[vkn_col].apply(
        lambda x: str(x).strip().replace('.0', '').replace(' ', ''))

    # Kaç VKN düzeltildi bilgisini logla
    duzeltilen = (df[vkn_col].apply(lambda x: len(str(x).strip().replace('.0','').replace(' ',''))).isin([8,9]) & mask_gecerli).sum()
    if duzeltilen > 0:
        log_cb(f"  🔧 {duzeltilen} satırda önde sıfır eksikti, otomatik tamamlandı (ör: 71419747 → 0071419747)", "warn")
    # Sağlam ayrıştırıcı: "1.234.567,89" gibi binlik ayraçlı tutarlar 0 olmasın
    df_t['_tutar'] = (df_t[tutar_col].apply(lambda v: para_deger(v) or 0.0)
                      if tutar_col else 0.0)

    # Geçersiz satır sayısını logla
    if len(df_gecersiz) > 0:
        log_cb(f"  ⚠️  {len(df_gecersiz)} geçersiz satır atlandı "
               f"(hatalı/eksik VKN-TC kimlik)", "warn")

    # ── %80 hedefi LİSTENİN TAMAMI üzerinden hesaplanır ──
    # Geçersiz VKN'li satırlar için tutanak üretilemez, ancak bu satırların
    # tutarları da listenin bir parçasıdır ve paydadan DÜŞÜLEMEZ. Aksi halde
    # gerçek kapsam %80'in altına düşer (yasal kural listenin bütünü içindir).
    TOPLAM_GECERLI = df_t['_tutar'].sum()
    TOPLAM_GECERSIZ = (df_gecersiz[tutar_col].apply(para_deger).fillna(0).sum()
                       if (tutar_col and len(df_gecersiz)) else 0.0)
    TOPLAM_LISTE = TOPLAM_GECERLI + TOPLAM_GECERSIZ
    HEDEF        = TOPLAM_LISTE * (yuzde80 / 100.0)
    # Boş liste / tümü sıfır tutar durumunda %oran hesaplarında bölme hatası olmasın
    pct = lambda x: (x / TOPLAM_LISTE * 100.0) if TOPLAM_LISTE else 0.0

    if TOPLAM_LISTE <= 0:
        log_cb("  ⚠️  Listede toplanabilir tutar bulunamadı "
               "(tutar sütunu boş ya da tüm satırlar geçersiz). "
               "Kaynak listeyi kontrol edin.", "warn")

    log_cb(f"  📊 Liste toplam: {TOPLAM_LISTE:,.2f} ₺  |  "
           f"%{yuzde80:.0f} hedef: {HEDEF:,.2f} ₺", "info")
    if TOPLAM_GECERSIZ > 0:
        log_cb(f"     (bunun {TOPLAM_GECERSIZ:,.2f} ₺'si geçersiz kimlikli satırlarda — "
               f"tutanak üretilemez ama %{yuzde80:.0f} hesabına dahildir)", "warn")

    # Firma bazında özet
    firma_ozet = {}
    for vkn, grp in df_t.groupby(vkn_col):
        uv = grp[unvan_col].dropna().unique() if unvan_col else []
        firma_ozet[vkn] = {
            'grp':    grp.drop(columns=['_tutar']),
            'toplam': grp['_tutar'].sum(),
            'max':    grp['_tutar'].max(),
            'unvan':  str(uv[0]) if len(uv) else vkn,
        }

    # ── 1. Aşama: 150K / 450K kriteri ──
    secilen   = {}
    kalan     = []

    for vkn, ozet in firma_ozet.items():
        if ozet['max'] >= esik_tek:
            secilen[vkn] = (ozet['grp'], f"tek fatura {ozet['max']:,.0f} ₺")
        elif ozet['toplam'] >= esik_toplam:
            secilen[vkn] = (ozet['grp'], f"toplam {ozet['toplam']:,.0f} ₺")
        else:
            kalan.append(vkn)

    secilen_tutar = sum(firma_ozet[v]['toplam'] for v in secilen)
    log_cb(f"  1️⃣  Kriter (tek≥{esik_tek:,.0f} / toplam≥{esik_toplam:,.0f}): "
           f"{len(secilen)} firma → {secilen_tutar:,.2f} ₺ "
           f"(%{pct(secilen_tutar):.1f})", "ok")

    # ── 2. Aşama: %80 tamamlama ──
    if secilen_tutar < HEDEF:
        kalan_sirali = sorted(kalan,
                              key=lambda v: firma_ozet[v]['toplam'],
                              reverse=True)
        ekstra = 0
        log_cb(f"  2️⃣  %{yuzde80:.0f} için ek firma ekleniyor:", "warn")
        for vkn in kalan_sirali:
            if secilen_tutar >= HEDEF:
                break
            ozet = firma_ozet[vkn]
            secilen[vkn] = (ozet['grp'],
                            f"%{yuzde80:.0f} tamamlama ({ozet['toplam']:,.0f} ₺)")
            secilen_tutar += ozet['toplam']
            ekstra += 1
            log_cb(f"     + {ozet['unvan'][:40]}  → "
                   f"kümülatif %{pct(secilen_tutar):.1f}", "warn")
        log_cb(f"  {ekstra} ek firma eklendi.", "warn")
        if secilen_tutar < HEDEF:
            log_cb(f"  ⛔ DİKKAT: Tüm uygun firmalar seçildiği halde %{yuzde80:.0f} "
                   f"hedefine ULAŞILAMADI (%{pct(secilen_tutar):.1f}). "
                   f"Sebep: geçersiz kimlikli satırlar ({TOPLAM_GECERSIZ:,.0f} ₺). "
                   f"Bu satırların VKN/TC bilgilerini listede düzeltip tekrar çalıştırın.", "err")
    else:
        log_cb(f"  ✅ 1. aşama zaten %{pct(secilen_tutar):.1f} — "
               f"ek firma gerekmedi.", "ok")

    log_cb(f"  📋 Toplam kapsam: {len(secilen)} firma  "
           f"(%{pct(secilen_tutar):.1f})", "ok")
    # Büyükten küçüğe sırala (dosya isimlendirme için)
    secilen_sirali = dict(
        sorted(secilen.items(),
               key=lambda kv: firma_ozet[kv[0]]['toplam'],
               reverse=True)
    )
    return secilen_sirali, df_gecersiz

# ══════════════════════════════════════════
#  EXCEL ÜRETME — Sistemin şablonu birebir
# ══════════════════════════════════════════
def guvenli_kaydet(wb, tam_yol):
    """Excel'i kaydeder. Windows'ta yol çok uzunsa (~260 karakter MAX_PATH sınırı)
    veya dosya adı geçersizse, adı kısaltarak yeniden dener; böylece uzun ünvanlı
    firmalarda dosya oluşturulamayıp numara atlaması yaşanmaz.
    Kaydedilen gerçek yolu döndürür."""
    yol    = Path(tam_yol)
    klasor = yol.parent
    stem   = yol.stem              # "N) 07_2026_VKN_UNVAN" (baştaki numara benzersizdir)
    # 1) Olduğu gibi dene
    try:
        wb.save(str(yol)); return str(yol)
    except Exception:
        pass
    # 2) Adı kısaltarak dene — baştaki "N) dönem_VKN" korunur, uzun ünvan kısalır
    for uzunluk in (80, 60, 40, 25, 15):
        try:
            kisa = stem[:uzunluk].rstrip(' ._-')
            yeni = klasor / f"{kisa}.xlsx"
            wb.save(str(yeni)); return str(yeni)
        except Exception:
            continue
    # 3) Windows uzun-yol ön eki ile son bir deneme
    if os.name == 'nt':
        try:
            uzun_pref = '\\\\?\\' + os.path.abspath(str(yol))
            wb.save(uzun_pref); return str(yol)
        except Exception:
            pass
    # Hiçbiri olmadıysa hatayı çağırana bildir
    wb.save(str(yol))
    return str(yol)


def firma_excel_olustur(firma_df, cikis_dosya, tum_kolonlar):
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = "Veriler"

    for ci, b in enumerate(SABLON_SUTUNLAR, 1):
        c = ws.cell(1, ci, value=b); c.number_format = '@'

    tarih_col   = sutun_bul(tum_kolonlar, ["alış faturasının tarihi"])
    # Fatura no: olduğu gibi (BBK2026... → fatura no)
    faturano_col= sutun_bul(tum_kolonlar, ["alış faturasının sıra no"])
    # Seri: yalnızca gerçek bir seri sütunu varsa; numara sütununu seri sanma
    seri_col    = seri_sutunu_bul(tum_kolonlar, tarih_col, faturano_col)
    tutar_col   = sutun_bul(tum_kolonlar, ["kdv hariç tutarı", "faturanın tutarı"])
    kdv_col     = kdv_sutunu_bul(tum_kolonlar)
    if not kdv_col: kdv_col = sutun_bul(tum_kolonlar, ["toplam indirilecek kdv"])
    defter_col  = sutun_bul(tum_kolonlar, ["defter kayıt tarihi"])
    yevmiye_col = sutun_bul(tum_kolonlar, ["yevmiye numarası"])
    odeme_col   = sutun_bul(tum_kolonlar, ["ödeme şekli"])
    # Açıklama: listedeki "Alınan Mal ve/veya Hizmetin Cinsi" sütununu kullan,
    # yoksa varsa bir "Açıklama" sütununa düş
    aciklama_col= sutun_bul(tum_kolonlar, ["alınan mal ve/veya hizmetin cinsi", "hizmetin cinsi", "cinsi"])
    if not aciklama_col:
        aciklama_col = sutun_bul(tum_kolonlar, ["açıklama"])

    for ri, (_, row) in enumerate(firma_df.iterrows()):
        er = ri + 2
        def yaz(ci, val, tip=''):
            # Para modu: gerçek SAYI olarak yaz (küsürat korunur, metin değil)
            if tip == 'para':
                f = para_deger(val)
                if f is None:
                    ws.cell(er, ci, value=None)
                else:
                    c = ws.cell(er, ci, value=f)
                    c.number_format = '0.00'
                return
            if   tip == 't': val = tarih_fmt(val)
            elif tip == 'n': val = sayi_fmt(val)
            else: val = '' if pd.isna(val) else str(val).strip()
            if val is None or val in ('None', 'nan'): val = ''
            c = ws.cell(er, ci, value=val); c.number_format = '@'

        # 1: Tarih
        yaz(1, row.get(tarih_col, None) if tarih_col else None, 't')

        # 2: Faturanın Serisi — ayrı seri sütunu varsa ve doluysa al, yoksa boş
        seri_val = ''
        if seri_col:
            sv = row.get(seri_col, None)
            if pd.notna(sv) and str(sv).strip() not in ('', 'None', 'nan'):
                seri_val = str(sv).strip()
        # Güvence: seri, fatura numarasıyla birebir aynıysa bu bir seri değildir → boş bırak
        if seri_val and faturano_col:
            fno_ayni = row.get(faturano_col, None)
            if pd.notna(fno_ayni) and str(fno_ayni).strip() == seri_val:
                seri_val = ''
        yaz(2, seri_val if seri_val else '')

        # 3: Faturanın Numarası — komple fatura no (BBK2026000001422)
        fno = row.get(faturano_col, None) if faturano_col else None
        yaz(3, '' if pd.isna(fno) else str(fno).strip())

        # 4-5: Tutar (matrah) ve KDV — gerçek SAYI olarak (küsürat korunur)
        yaz(4, row.get(tutar_col,    None) if tutar_col    else None, 'para')
        yaz(5, row.get(kdv_col,      None) if kdv_col      else None, 'para')

        # 6-10
        yaz(6, row.get(defter_col,   None) if defter_col   else None, 't')
        yaz(7, row.get(yevmiye_col,  None) if yevmiye_col  else None)
        yaz(8, row.get(odeme_col,    None) if odeme_col    else None)
        yaz(9, row.get(aciklama_col, None) if aciklama_col else None)
        yaz(10, '')

    return guvenli_kaydet(wb, cikis_dosya)

# ══════════════════════════════════════════
#  PDF ÇIKTI (opsiyonel — reportlab varsa)
# ══════════════════════════════════════════
def _pdf_font_bul():
    """Türkçe karakterleri düzgün gösteren bir TTF'yi reportlab'a kaydeder;
    bulunursa kayıtlı font adını, bulunamazsa 'Helvetica' döndürür (Helvetica
    bazı Türkçe karakterleri bozabilir, bu yüzden Unicode TTF tercih edilir)."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception:
        return 'Helvetica'
    adaylar = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/dejavu/DejaVuSans.ttf',
        'C:\\Windows\\Fonts\\arial.ttf',
        'C:\\Windows\\Fonts\\segoeui.ttf',
        '/Library/Fonts/Arial.ttf',
        '/System/Library/Fonts/Supplemental/Arial.ttf',
    ]
    for y in adaylar:
        try:
            if os.path.exists(y):
                pdfmetrics.registerFont(TTFont('TutanakFont', y))
                return 'TutanakFont'
        except Exception:
            continue
    return 'Helvetica'

def pdf_destekli():
    """reportlab kurulu mu? (PDF üretimi opsiyoneldir)"""
    try:
        import reportlab  # noqa: F401
        return True
    except Exception:
        return False

def firma_pdf_olustur(firma_df, pdf_dosya, tum_kolonlar, vkn='', unvan='', donem=''):
    """Bir firmanın tutanağının okunur bir PDF kopyasını üretir (arşiv/imza için).
    Resmî yükleme dosyası Excel'dir; PDF yalnızca insan-okur kopyadır.
    reportlab kurulu değilse ImportError yükseltir (çağıran atlar)."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    font = _pdf_font_bul()
    tarih_col    = sutun_bul(tum_kolonlar, ["alış faturasının tarihi"])
    faturano_col = sutun_bul(tum_kolonlar, ["alış faturasının sıra no"])
    seri_col     = seri_sutunu_bul(tum_kolonlar, tarih_col, faturano_col)
    tutar_col    = sutun_bul(tum_kolonlar, ["kdv hariç tutarı", "faturanın tutarı"])
    kdv_col      = kdv_sutunu_bul(tum_kolonlar) or sutun_bul(tum_kolonlar, ["toplam indirilecek kdv"])
    aciklama_col = sutun_bul(tum_kolonlar, ["alınan mal ve/veya hizmetin cinsi", "hizmetin cinsi", "cinsi"]) \
                   or sutun_bul(tum_kolonlar, ["açıklama"])

    def _para_str(v):
        f = para_deger(v)
        return '' if f is None else f'{f:,.2f}'.replace(',', '¤').replace('.', ',').replace('¤', '.')

    basliklar = ["Tarih", "Seri", "Fatura No", "Matrah (TL)", "KDV (TL)", "Açıklama"]
    satirlar = [basliklar]
    top_matrah = top_kdv = 0.0
    for _, row in firma_df.iterrows():
        seri_val = ''
        if seri_col:
            sv = row.get(seri_col, None)
            if pd.notna(sv) and str(sv).strip() not in ('', 'None', 'nan'):
                seri_val = str(sv).strip()
        if seri_val and faturano_col:
            fno_ayni = row.get(faturano_col, None)
            if pd.notna(fno_ayni) and str(fno_ayni).strip() == seri_val:
                seri_val = ''
        mv = para_deger(row.get(tutar_col, None)) if tutar_col else None
        kv = para_deger(row.get(kdv_col, None)) if kdv_col else None
        top_matrah += mv or 0.0
        top_kdv    += kv or 0.0
        fno = row.get(faturano_col, None) if faturano_col else None
        acik = row.get(aciklama_col, None) if aciklama_col else None
        satirlar.append([
            tarih_fmt(row.get(tarih_col, None)) if tarih_col else '',
            seri_val,
            '' if pd.isna(fno) else str(fno).strip(),
            _para_str(mv), _para_str(kv),
            '' if (acik is None or pd.isna(acik)) else str(acik).strip()[:40],
        ])
    satirlar.append(["", "", "TOPLAM", _para_str(top_matrah), _para_str(top_kdv), ""])

    styles = getSampleStyleSheet()
    bas_st = ParagraphStyle('bas', parent=styles['Title'], fontName=font, fontSize=13)
    alt_st = ParagraphStyle('alt', parent=styles['Normal'], fontName=font, fontSize=9)
    dip_st = ParagraphStyle('dip', parent=styles['Normal'], fontName=font,
                            fontSize=7, textColor=colors.grey)

    doc = SimpleDocTemplate(pdf_dosya, pagesize=landscape(A4),
                            topMargin=15*mm, bottomMargin=12*mm,
                            leftMargin=12*mm, rightMargin=12*mm)
    icerik = [
        Paragraph("KARŞIT İNCELEME TUTANAĞI (okunur kopya)", bas_st),
        Spacer(1, 4*mm),
        Paragraph(f"<b>Ünvan:</b> {unvan or '-'} &nbsp;&nbsp; "
                  f"<b>VKN/TC:</b> {vkn or '-'} &nbsp;&nbsp; "
                  f"<b>Dönem:</b> {donem or '-'} &nbsp;&nbsp; "
                  f"<b>Fatura sayısı:</b> {len(firma_df)}", alt_st),
        Spacer(1, 4*mm),
    ]
    genislik = [26*mm, 16*mm, 42*mm, 34*mm, 30*mm, 90*mm]
    t = Table(satirlar, colWidths=genislik, repeatRows=1)
    t.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font),
        ('FONTSIZE', (0, 0), (-1, -1), 7.5),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A1A2E')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#EDEFF5')),
        ('FONTNAME', (0, -1), (-1, -1), font),
        ('ALIGN', (3, 0), (4, -1), 'RIGHT'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#B8BEC9')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F7F8FB')]),
    ]))
    icerik.append(t)
    icerik.append(Spacer(1, 6*mm))
    icerik.append(Paragraph(
        "Not: Resmî yükleme dosyası ilgili Excel tutanağıdır; bu PDF yalnızca "
        "okunur/arşiv kopyasıdır.", dip_st))
    doc.build(icerik)
    return pdf_dosya

# ══════════════════════════════════════════
#  WORD ŞABLON EŞLEŞTİRME (VKN ile) VE ÜRETME
#  Kullanıcının hazır .doc tutanak şablonları VKN'ye göre eşleştirilir; eşleşen
#  şablonun yalnızca "Karşıt İncelemeye Konu Fatura" tablosu, seçilen firmanın
#  faturalarıyla güncellenir. Diğer her şey (YMM, iade talep eden firma, defter
#  onayları) değişmeden kalır.
# ══════════════════════════════════════════
def _ascii_kucuk(s):
    """Türkçe karakterleri ASCII'ye indirip küçük harfe çevirir. Python'un
    str.lower()'ı Türkçe'de yanıltıcıdır ('İ'.lower() → noktalı 'i̇'), bu yüzden
    anahtar-kelime eşleşmelerinde bunu kullan (İNCELEME → inceleme)."""
    tr = {'İ': 'i', 'I': 'i', 'ı': 'i', 'Ş': 's', 'ş': 's', 'Ğ': 'g', 'ğ': 'g',
          'Ü': 'u', 'ü': 'u', 'Ö': 'o', 'ö': 'o', 'Ç': 'c', 'ç': 'c'}
    return ''.join(tr.get(c, c) for c in str(s)).lower()

def _tr_para_str(v):
    """Sayıyı Türkçe biçimde metne çevirir: 683200.0 → '683.200,00'. Boşsa ''."""
    f = para_deger(v)
    if f is None:
        return ''
    return f'{f:,.2f}'.replace(',', '¤').replace('.', ',').replace('¤', '.')

def _doc_metni_oku(path):
    """Eski ikili .doc (Word 97-2003) dosyasının ana metnini çıkarır.
    Metin WordDocument akışında UTF-16LE saklanır; hücre/satır sonları (0x07)
    sekmeye çevrilir. Yalnızca OKUMA içindir (VKN eşleştirme). `olefile` gerekir."""
    import olefile
    with olefile.OleFileIO(path) as ole:
        data = ole.openstream('WordDocument').read()
    t = data.decode('utf-16-le', 'ignore').replace('\r', '\n').replace('\x07', '\t')
    return re.sub(r'[\x00-\x06\x08\x0b\x0c\x0e-\x1f]', '', t)

def _docx_metni_oku(path):
    """Modern .docx dosyasının metnini çıkarır (paragraflar + tablo hücreleri).
    Her tablo satırının hücreleri sekme ile birleştirilir; böylece
    sablon_vkn_metinden'in etiket→değer (Ünvanı \\t X) örüntüsü .docx'te de
    çalışır. `python-docx` gerekir (yalnızca okuma)."""
    import docx
    doc = docx.Document(path)
    parcalar = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            parcalar.append('\t'.join(c.text.strip() for c in row.cells))
    return '\n'.join(parcalar)

def docx_destekli():
    """python-docx kullanılabilir mi? (.docx şablon okuma/yazma için)"""
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False

def _vkn_metinden_ayikla(s):
    """Bir metin parçasından geçerli VKN/TCKN (10-11 hane) çıkarır; yoksa None.
    Aradaki boşlukları temizler (ör. '493 061 9102' → '4930619102'), 8-9 haneyi
    önde sıfırla tamamlar; yer tutucu (tek-rakam) kimlikleri geçersiz sayar.
    firmalari_filtrele'deki normalize ile aynı kuralı kullanır ki VKN'ler eşleşsin."""
    en_iyi = None
    for parca in re.findall(r'\d(?:[\d ]*\d)?', str(s)):
        d = re.sub(r'\D', '', parca)
        if len(d) in (8, 9):
            d = d.zfill(10)
        if len(d) in (10, 11) and d != d[0] * len(d):
            en_iyi = d   # blok içindeki son geçerli kimlik (karşı firma hücresi)
    return en_iyi

def _blok_vkn(blok):
    """Karşı firma bloğundan VKN/TCKN ayıklar. Telefonla karışmaması için önce
    'V.D.' / 'Vergi Dairesi' / 'Kurumlar V.D.' civarındaki numarayı arar (bazı
    şablonlarda etiketler karışık olduğundan bu daha güvenilir); bulamazsa son
    çare olarak bloktaki geçerli ilk kimliği alır."""
    for m in re.finditer(
            r'(?:[Vv]\.?\s*[Dd]\.?|[Vv]ergi\s*[Dd]airesi|[Kk]urumlar)'
            r'\s*[–\-/:.]*\s*([0-9][0-9 ]{7,14}[0-9])', blok):
        n = _vkn_metinden_ayikla(m.group(1))
        if n:
            return n
    # Son çare: telefon/faks satırlarını dışla, kalan ilk geçerli kimliği al
    for satir in blok.split('\n'):
        if re.search(r'tel|faks|fax', satir, re.I):
            continue
        n = _vkn_metinden_ayikla(satir)
        if n:
            return n
    return None

def sablon_vkn_metinden(metin):
    """Bir tutanak/yazı metninden karşı firmanın (vkn, unvan) bilgisini çıkarır.
    İki belge tipini de tanır:
      • Karşıt inceleme tutanağı → 'NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN'
      • YMM Bilgi İsteme yazısı  → 'Hakkında Bilgi İstenilen Mükellef…'
    Bulunamazsa (None, None)."""
    m = re.search(
        r'(?:NEZD[İIı]NDE\s+KAR[ŞSş]IT\s+[İIı]NCELEME\s+YAPILAN\s+F[İIı]RMANIN'
        r'|HAKKINDA\s+B[İIı]LG[İIı]\s+[İIı]STEN[İIı]LEN\s+M[ÜUü]KELLEF[İIı]?[^\t\n]*)'
        r'(.*?)'
        r'(?:[İIı]NCELEME\s+DAYANA|[İIı]STEN[İIı]LEN\s+B[İIı]LG[İIı]LER'
        r'|Kar[şs]ıt\s+İncelemeye|$)',
        metin, re.S | re.I)
    if not m:
        return (None, None)
    blok = m.group(1)
    unv = re.search(r'Ünvan[ıi]\s*\t+([^\t\n]+)', blok)
    if not unv:
        unv = re.search(r'[ÜUü]nvan[ıi]\s*\t+([^\t\n]+)', blok)
    unvan = unv.group(1).strip() if unv else None
    # Önce doğru etiketli hücre; olmazsa 'V.D.' çevresinden ayıkla (karışık etiket)
    vd = re.search(r'Vergi\s*Dairesi\s*/?\s*Nosu\s*\t+([^\t\n]+)', blok)
    vkn = _vkn_metinden_ayikla(vd.group(1)) if vd else None
    if not vkn:
        vkn = _blok_vkn(blok)
    return (vkn, unvan)

def sablon_vkn_oku(path):
    """Bir .doc/.docx şablonundan (vkn, unvan) döndürür; okunamaz/eşleşmezse
    (None, None). Uzantıya göre doğru okuyucuyu seçer. (Tek firmalık dosyalar için;
    çok firmalı .docx için _sablon_kayitlari kullanın.)"""
    try:
        if str(path).lower().endswith('.docx'):
            metin = _docx_metni_oku(path)
        else:
            metin = _doc_metni_oku(path)
        return sablon_vkn_metinden(metin)
    except Exception:
        return (None, None)

def _metni_bloklara_ayir(metin):
    """Bir tutanak METNİNİ 'KATMA DEĞER…KARŞIT İNCELEME…TUTANAĞI' başlıklarından
    bloklara ayırır; her blok bir firma tutanağıdır. Tek/başlıksız metinde tek
    eleman döner. (Eski ikili .doc'un birleşik olup olmadığını anlamakta kullanılır.)"""
    ak = _ascii_kucuk(metin)
    # İki belge tipinin de blok başlangıcı: (a) karşıt inceleme TUTANAĞI başlığı,
    # (b) YMM (Bilgi İsteme) YAZISI'nın 'Konu : Bilgi İsteme' satırı.
    idxs = sorted(m.start() for m in re.finditer(
        r'katma deger.{0,40}?karsit inceleme.{0,20}?tutana|konu\s*[:：].{0,20}?bilgi isteme',
        ak))
    if len(idxs) <= 1:
        return [metin]
    idxs.append(len(metin))
    return [metin[idxs[i]:idxs[i + 1]] for i in range(len(idxs) - 1)]

def _doc_docx_cevir(path):
    """Eski ikili .doc'u Word (COM) ile geçici bir .docx'e çevirir; yolu döndürür.
    Yalnızca Windows + Word'de çalışır (birleşik .doc'ları bölmek için gerekir);
    pywin32/Word yoksa RuntimeError yükseltir. Şablonun aslı değişmez."""
    try:
        import win32com.client as win32
    except Exception:
        raise RuntimeError("Word otomasyonu yok (birleşik .doc'u bölmek için gerekli)")
    import tempfile
    hedef = os.path.join(tempfile.mkdtemp(prefix='exay_'), Path(path).stem + '.docx')
    word = win32.DispatchEx("Word.Application"); word.Visible = False
    try:
        try: word.DisplayAlerts = 0
        except Exception: pass
        d = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
        d.SaveAs(os.path.abspath(hedef), FileFormat=16)   # 16 = wdFormatDocumentDefault (.docx)
        d.Close(SaveChanges=False)
    finally:
        try: word.Quit()
        except Exception: pass
    return hedef

def _sablon_kayitlari(path):
    """Bir şablon dosyasındaki TÜM firma kayıtlarını [(vkn, unvan, blok, üretim_yolu)]
    döndürür. Çok firmalı .docx → her blok ayrı (blok=indeks). Tek .doc → (blok=None).
    Birleşik .doc (birden çok tutanak) → Word ile .docx'e çevrilir ve bloklar oradan
    okunur (üretim_yolu geçici .docx olur; üretim böylece test edilmiş .docx yolundan gider)."""
    p = str(path)
    try:
        if p.lower().endswith('.docx'):
            import docx
            bloklar = _docx_firma_bloklari(docx.Document(p))
            return [(b['vkn'], b['unvan'], k, p)
                    for k, b in enumerate(bloklar) if b['vkn']]
        # .doc
        metin = _doc_metni_oku(p)
        if len(_metni_bloklara_ayir(metin)) <= 1:
            vkn, unvan = sablon_vkn_metinden(metin)
            return [(vkn, unvan, None, p)] if vkn else []
        # Birleşik .doc → .docx'e çevir, blokları oradan indeksle
        dx = _doc_docx_cevir(p)
        import docx
        bloklar = _docx_firma_bloklari(docx.Document(dx))
        return [(b['vkn'], b['unvan'], k, dx)
                for k, b in enumerate(bloklar) if b['vkn']]
    except Exception:
        return []

def sablonlari_indeksle(klasor, log_cb=None):
    """Verilen klasördeki (alt klasörler dahil) .doc VE .docx şablonlarını karşı
    firma VKN'sine göre indeksler → {vkn: dosya_yolu}. Aynı VKN birden çok
    şablonda varsa sonuncusu kullanılır ve uyarılır. Klasör yoksa boş sözlük."""
    idx = {}
    kok = Path(klasor) if klasor else None
    if not kok or not kok.exists():
        return idx
    dosyalar = sorted(list(kok.rglob('*.doc')) + list(kok.rglob('*.docx')))
    var_doc  = any(p.suffix.lower() == '.doc' for p in dosyalar)
    var_docx = any(p.suffix.lower() == '.docx' for p in dosyalar)
    # Okuma için gereken kütüphaneler yoksa net uyar
    if var_doc:
        try:
            import olefile  # noqa: F401
        except Exception:
            if log_cb:
                log_cb("  ⚠️  'olefile' kurulu değil; .doc şablonlar okunamaz "
                       "(pip install olefile).", "warn")
    if var_docx and not docx_destekli() and log_cb:
        log_cb("  ⚠️  'python-docx' kurulu değil; .docx şablonlar okunamaz "
               "(pip install python-docx).", "warn")
    coklu = 0
    for p in dosyalar:
        if p.name.startswith('~$'):        # Word geçici dosyaları
            continue
        kayitlar = _sablon_kayitlari(str(p))
        if len(kayitlar) > 1:
            coklu += 1
        for vkn, unvan, blok, yol in kayitlar:
            if vkn in idx and log_cb:
                log_cb(f"  ⚠️  Aynı VKN ({vkn}) için birden çok kayıt; sonuncusu "
                       f"kullanılacak: {p.name}", "warn")
            idx[vkn] = (yol, blok)
    if coklu and log_cb:
        log_cb(f"  🧩 {coklu} dosya çok-firmalı (tek Word'de birden çok tutanak) "
               f"olarak tanındı; her firma ayrı ayrı eşleştirildi.", "info")
    return idx

def word_destekli():
    """Word otomasyonu (pywin32 + Windows Word) kullanılabilir mi?"""
    try:
        import win32com.client  # noqa: F401
        return True
    except Exception:
        return False

def _word_fatura_satiri(row, kolonlar):
    """Bir fatura satırını Word tablosu sırasına göre hücre listesine çevirir:
    [Tarih, Fatura No, Malın Cinsi, Miktarı, Tutarı, KDV Tutarı, Defter Kayıt].
    Kaynak sütunları exay'ın standart bulucularıyla eşlenir."""
    tarih_col    = sutun_bul(kolonlar, ["alış faturasının tarihi", "fatura tarihi", "tarih"])
    faturano_col = sutun_bul(kolonlar, ["alış faturasının sıra no", "fatura no"])
    tutar_col    = sutun_bul(kolonlar, ["kdv hariç tutarı", "faturanın tutarı"])
    kdv_col      = kdv_sutunu_bul(kolonlar) or sutun_bul(kolonlar, ["toplam indirilecek kdv"])
    cinsi_col    = sutun_bul(kolonlar, ["alınan mal ve/veya hizmetin cinsi", "hizmetin cinsi", "cinsi"]) \
                   or sutun_bul(kolonlar, ["açıklama"])
    miktar_col   = sutun_bul(kolonlar, ["miktar"])
    def al(c):
        if not c: return ''
        v = row.get(c, None)
        return '' if (v is None or pd.isna(v)) else str(v).strip()
    return [
        tarih_fmt(row.get(tarih_col, None)) if tarih_col else '',
        al(faturano_col),
        al(cinsi_col),
        al(miktar_col),
        _tr_para_str(row.get(tutar_col, None)) if tutar_col else '',
        _tr_para_str(row.get(kdv_col, None)) if kdv_col else '',
        '',   # Defter Kayıt Tarihi/Nosu — kullanıcı isteğiyle boş
    ]

# Fatura tablosunda 'başlık satırı' sayılacak anahtarlar (ASCII, küçük harf).
# _ascii_kucuk ile karşılaştırılır ki Türkçe İ/ı casing sorun çıkarmasın.
_BAS_KELIME = ('faturanin', 'malin', 'tarih', 'numar', 'cins',
               'miktar', 'tutar', 'kdv', 'defter', 'nosu', 'belge')

def _fatura_tablosu_mu(basliklar):
    """Bir Word tablosunun başlık hücre metinlerine bakıp 'Karşıt İncelemeye Konu
    Fatura' tablosu olup olmadığına karar verir."""
    b = _ascii_kucuk(' '.join(basliklar))
    # Numara sütunu tutanakta 'Numarası', YMM (Bilgi İsteme) yazısında 'Nosu'.
    numar = ('numar' in b) or ('nosu' in b) or ('fat. no' in b) or ('fat.no' in b)
    imzalar = int('tarih' in b) + int(numar) + int('kdv' in b)
    return imzalar >= 2 and ('cins' in b or 'tutar' in b or 'matrah' in b)

def _fatura_kolon_rolu(baslik):
    """Bir fatura sütununun başlık metninden rolünü belirler. İki belge tipini de
    (karşıt inceleme TUTANAĞI ve YMM 'Bilgi İsteme' YAZISI) tek eşlemeyle karşılar.
    Roller: tarih, no, cins, miktar, matrah, kdv, dahil (KDV dahil toplam), bos."""
    h = _ascii_kucuk(baslik)
    if 'defter' in h or 'yevmiye' in h or 'kayit' in h:
        return 'bos'                       # tutanaktaki 'Defter Kayıt' sütunu → boş
    if 'tarih' in h:
        return 'tarih'
    if 'numar' in h or 'nosu' in h or ('no' in h and 'cins' not in h and 'kdv' not in h):
        return 'no'
    if 'cins' in h:
        return 'cins'
    if 'miktar' in h:
        return 'miktar'
    if 'dahil' in h or 'toplam' in h:
        return 'dahil'                     # YMM yazısındaki 'KDV dahil toplam'
    if 'kdv' in h:
        return 'kdv'
    if 'matrah' in h or 'tutar' in h:
        return 'matrah'
    return 'bos'

def _fatura_deger_haritasi(row, kolonlar):
    """Kaynak fatura satırından her rol için hücre metnini üretir (rol → metin).
    matrah/kdv aynı bulucularla; 'dahil' = matrah + kdv (sayısal, sonra TR biçim)."""
    tarih, no, cins, miktar, matrah, kdv, _bos = _word_fatura_satiri(row, kolonlar)
    tutar_col = sutun_bul(kolonlar, ["kdv hariç tutarı", "faturanın tutarı"])
    kdv_col   = kdv_sutunu_bul(kolonlar) or sutun_bul(kolonlar, ["toplam indirilecek kdv"])
    dahil = ''
    try:
        mv = para_deger(row.get(tutar_col, None)) if tutar_col else None
        kv = para_deger(row.get(kdv_col, None)) if kdv_col else None
        if mv is not None or kv is not None:
            dahil = _tr_para_str((mv or 0.0) + (kv or 0.0))
    except Exception:
        dahil = ''
    return {'tarih': tarih, 'no': no, 'cins': cins, 'miktar': miktar,
            'matrah': matrah, 'kdv': kdv, 'dahil': dahil, 'bos': ''}

def _fatura_kolon_basliklari(satirlar, veri_bas, hucre_metni=lambda c: c.text):
    """Fatura tablosunun her sütunu için, TÜM başlık satırlarını (0..veri_bas-1)
    birleştiren başlık metnini döndürür. Birleşik (merged) hücreler python-docx'te
    sütun boyunca tekrar ettiğinden sütun indeksiyle güvenle hizalanır. Böylece
    tutanaktaki iki satırlı 'Defter Kayıt / Tarihi-Nosu' başlığı tek metinde toplanır."""
    bas_satirlar = list(satirlar)[:max(veri_bas, 1)]
    if not bas_satirlar:
        return []
    n = max(len(getattr(r, 'cells', r)) for r in bas_satirlar)
    basliklar = [''] * n
    for r in bas_satirlar:
        for ci, c in enumerate(getattr(r, 'cells', r)):
            if ci < n:
                basliklar[ci] = (basliklar[ci] + ' ' + hucre_metni(c)).strip()
    return basliklar

def _fatura_satir_hucreleri(row, kolonlar, roller):
    """Rol listesine göre bir kaynak satırı hedef tablo sütun sırasına çevirir.
    `roller` güvenilir görünmüyorsa (tarih ve para sütunu yoksa) None döner ki
    çağıran eski KONUMSAL eşlemeye (`_word_fatura_satiri`) düşebilsin."""
    if not roller or 'tarih' not in roller or not ({'matrah', 'kdv', 'dahil'} & set(roller)):
        return None
    harita = _fatura_deger_haritasi(row, kolonlar)
    return [harita.get(rol, '') for rol in roller]

def firma_word_olustur(sablon_yol, firma_df, cikis_yol, tum_kolonlar, log_cb=None,
                       inceleme_dayanagi=None):
    """Eşleşen .doc şablonunu Word (COM) ile açar, 'Karşıt İncelemeye Konu Fatura'
    tablosunun VERİ satırlarını firma faturalarıyla değiştirir ve `cikis_yol`'a
    YENİ dosya olarak kaydeder. Şablonun kendisi değiştirilmez.

    Yalnızca Windows + Word kuruluyken (pywin32) çalışır; aksi halde RuntimeError
    yükseltir ve çağıran zarifçe atlar (Excel çıktısı yine üretilir).

    NOT: Word tablosunun sütun düzeni/satır sayısı gerçek şablonlara göre burada
    (Word olmayan ortamda) doğrulanamadığından, bu adım kullanıcı makinesinde
    test edilmeli; log ayrıntısı (tablo bulundu, sütun/satır sayısı) buna göre
    tutulur."""
    try:
        import win32com.client as win32
        from win32com.client import constants as C  # noqa: F401
    except Exception:
        raise RuntimeError("Word otomasyonu yok (pywin32 + Windows Word gerekli)")

    def _yaz(m, t='info'):
        if log_cb: log_cb(m, t)

    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    try:
        word.DisplayAlerts = 0
    except Exception:
        pass
    # Bir hücrenin temiz metni (hücre/paragraf işaretleri atılır). Word COM'da
    # birleşik (merged) hücreli tablolarda Columns.Count hata verebildiği için,
    # her yerde satırın kendi Cells koleksiyonu üzerinden gidilir.
    def _hucre(cell):
        try:
            return re.sub(r'[\r\x07\x02\n]', '', cell.Range.Text).strip()
        except Exception:
            return ''
    def _satir_metni(tbl, r):
        try:
            return ' '.join(_hucre(c) for c in tbl.Rows(r).Cells)
        except Exception:
            return ''

    doc = None
    try:
        doc = word.Documents.Open(os.path.abspath(sablon_yol), ReadOnly=False)
        hedef_tablo = None
        veri_bas = 2
        for tbl in doc.Tables:
            try:
                rmax = tbl.Rows.Count
            except Exception:
                continue
            # İlk (en çok) iki satırın metnini başlık adayı olarak topla
            basliklar = []
            for r in range(1, min(2, rmax) + 1):
                try:
                    basliklar.extend(_hucre(c) for c in tbl.Rows(r).Cells)
                except Exception:
                    pass
            if not _fatura_tablosu_mu(basliklar):
                continue
            hedef_tablo = tbl
            # Veri başlangıcı: baştan başlık gibi görünen satırları (FATURANIN,
            # MALIN, Tarihi, Numarası, Cinsi, Miktarı, Tutarı, KDV, Defter...)
            # atla; ilk veri satırında dur.
            veri_bas = 1
            for r in range(1, rmax + 1):
                txt = _ascii_kucuk(_satir_metni(tbl, r))
                if txt.strip() and any(k in txt for k in _BAS_KELIME):
                    veri_bas = r + 1
                else:
                    break
            break
        if hedef_tablo is None:
            raise RuntimeError("Şablonda fatura tablosu bulunamadı")

        _yaz(f"      🧩 Fatura tablosu bulundu (veri {veri_bas}. satırdan başlıyor).", "info")

        # Sütun rollerini başlık satırlarından (1..veri_bas-1) çöz. Table.Cell(r,c)
        # birleşik hücrelerde bile ilgili sütunun hücresini döndürdüğünden, iki
        # satırlı 'Defter Kayıt / Tarihi-Nosu' başlığı tek metinde birleşir ve
        # tutanak (son sütun boş) ile YMM yazısı (son sütun KDV dahil toplam)
        # doğru ayrışır. Çözülemezse KONUMSAL eşlemeye düşülür.
        roller = None
        try:
            ncol = hedef_tablo.Rows(max(veri_bas - 1, 1)).Cells.Count
            kol_bas = []
            for c in range(1, ncol + 1):
                parc = []
                for r in range(1, veri_bas):
                    try:
                        parc.append(_hucre(hedef_tablo.Cell(r, c)))
                    except Exception:
                        pass
                kol_bas.append(' '.join(parc))
            roller = [_fatura_kolon_rolu(b) for b in kol_bas]
        except Exception:
            roller = None

        # Mevcut veri satırlarını sil (baştaki başlık satırları korunur)
        while hedef_tablo.Rows.Count >= veri_bas:
            try:
                hedef_tablo.Rows(hedef_tablo.Rows.Count).Delete()
            except Exception:
                break

        yazilan = 0
        for _, row in firma_df.iterrows():
            hucreler = _fatura_satir_hucreleri(row, tum_kolonlar, roller) \
                       or _word_fatura_satiri(row, tum_kolonlar)
            yeni = hedef_tablo.Rows.Add()
            try:
                cells = yeni.Cells
                n = cells.Count
            except Exception:
                n = len(hucreler)
                cells = None
            for ci in range(1, n + 1):
                if ci - 1 < len(hucreler):
                    try:
                        (cells(ci) if cells else yeni.Cells(ci)).Range.Text = hucreler[ci - 1]
                    except Exception:
                        pass
            yazilan += 1

        # İnceleme Dayanağı (sözleşme) verildiyse etiketin yanındaki hücreyi güncelle
        if inceleme_dayanagi:
            yazildi = False
            for tbl in doc.Tables:
                try:
                    rmax = tbl.Rows.Count
                except Exception:
                    continue
                for r in range(1, rmax + 1):
                    try:
                        hucreler2 = list(tbl.Rows(r).Cells)
                    except Exception:
                        continue
                    for ci, hc in enumerate(hucreler2):
                        et = _ascii_kucuk(re.sub(r'[\r\x07\x02]', '', hc.Range.Text))
                        if 'inceleme dayana' in et and ci + 1 < len(hucreler2):
                            try:
                                hucreler2[ci + 1].Range.Text = inceleme_dayanagi
                                yazildi = True
                            except Exception:
                                pass
                            break
                    if yazildi:
                        break
                if yazildi:
                    break
            if not yazildi:
                _yaz("      ⚠️ 'İNCELEME DAYANAĞI' alanı şablonda bulunamadı; sözleşme "
                     "bilgisi güncellenemedi.", "warn")

        os.makedirs(os.path.dirname(os.path.abspath(cikis_yol)), exist_ok=True)
        # 0 = wdFormatDocument (.doc). Şablonla aynı biçimde kaydet.
        doc.SaveAs(os.path.abspath(cikis_yol), FileFormat=0)
        _yaz(f"      📝 Word tutanağı: {Path(cikis_yol).name} ({yazilan} fatura satırı)", "ok")
        return cikis_yol
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            word.Quit()
        except Exception:
            pass

def _guvenli_docx_kaydet(doc, tam_yol):
    """.docx'i kaydeder; Windows uzun yol/uzun dosya adı sorununda adı kısaltarak
    yeniden dener (guvenli_kaydet'in .docx karşılığı). Kaydedilen gerçek yolu döndürür."""
    yol = Path(tam_yol); klasor = yol.parent; stem = yol.stem
    try:
        doc.save(str(yol)); return str(yol)
    except Exception:
        pass
    for uzunluk in (80, 60, 40, 25, 15):
        try:
            kisa = stem[:uzunluk].rstrip(' ._-')
            yeni = klasor / f"{kisa}.docx"
            doc.save(str(yeni)); return str(yeni)
        except Exception:
            continue
    doc.save(str(yol)); return str(yol)   # son deneme (hatayı çağırana bildirir)

def _docx_hucre_yaz(cell, metin):
    """Hücreye metni yazar; mevcut paragraf biçimini (hizalama/stil) ve varsa ilk
    run'ın font biçimini KORUYARAK. Böylece şablonun hücre düzeni (ör. ortalı/sağa
    dayalı, özel font) bozulmadan kalır. `cell.text = ...` bunları sıfırlardı."""
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:      # tek paragraf kalsın
        extra._p.getparent().remove(extra._p)
    if p.runs:
        p.runs[0].text = metin             # ilk run'ı kullan (font korunur)
        for r in p.runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        p.add_run(metin)

def _docx_nezdinde_yaz(doc, unvan, vkn):
    """Boş/yedek şablonda 'NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN' bloğundaki
    Ünvanı ve Vergi Dairesi/Nosu değer hücrelerine bilinen ünvan/VKN'yi yazar
    (listeden gelir). Yalnızca NEZDİNDE bölümünü hedefler; İADE TALEBİNDE bloğuna
    dokunmaz. En az biri yazıldıysa True."""
    for tbl in doc.tables:
        rows = tbl.rows
        nez_i = None
        for i, r in enumerate(rows):
            if r.cells and 'nezdinde' in _ascii_kucuk(r.cells[0].text):
                nez_i = i
                break
        if nez_i is None:
            continue
        yazildi = False
        for r in rows[nez_i + 1:]:
            et = _ascii_kucuk(r.cells[0].text) if r.cells else ''
            if 'inceleme dayana' in et or 'iade talebinde' in et:
                break
            if len(r.cells) < 2 or r.cells[1]._tc is r.cells[0]._tc:
                continue
            if unvan and et.strip() == 'unvani':
                _docx_hucre_yaz(r.cells[1], str(unvan)); yazildi = True
            elif 'vergi dairesi' in et:
                _docx_hucre_yaz(r.cells[1], str(vkn)); yazildi = True
        return yazildi
    return False

def _docx_inceleme_dayanagi_yaz(doc, metin):
    """'İNCELEME DAYANAĞI' satırının DEĞER hücresini `metin` ile günceller (tarih/
    sözleşme her yıl değişir; eski şablonun eski bilgisi otomatik ezilir).
    Bulup yazarsa True. Etiket hücresi ('İNCELEME DAYANAĞI') değişmez."""
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if cells and 'inceleme dayana' in _ascii_kucuk(cells[0].text):
                for c in cells[1:]:
                    if c._tc is not cells[0]._tc:   # birleşik değil, gerçek değer hücresi
                        _docx_hucre_yaz(c, metin)
                        return True
    return False

def _docx_firma_bloklari(doc):
    """Bir .docx gövdesini firma bloklarına ayırır. Her tutanak bloğu
    'KATMA DEĞER ... KARŞIT İNCELEME TUTANAĞI' başlık paragrafıyla başlar.
    [{'ilk':i,'son':j,'vkn':..,'unvan':..,'els':[...]}...] döndürür (els: bloğun
    gövde elemanları). Başlık yoksa tüm gövde tek blok sayılır (klasik şablon)."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    els = [el for el in doc.element.body if el.tag in (qn('w:p'), qn('w:tbl'))]
    bas = []
    for i, el in enumerate(els):
        if el.tag == qn('w:p'):
            t = _ascii_kucuk(Paragraph(el, doc).text)
            # Blok başı: karşıt inceleme TUTANAĞI başlığı (belge tepesinde) VEYA
            # YMM yazısının 'Konu : Bilgi İsteme' satırı. YMM yazısında bu satırın
            # ÜSTÜNDE 'Sayı :' başlık satırı vardır; blok başını ona geri çekeriz
            # ki üst başlık (Sayı/tarih) bloktan düşüp silinmesin.
            if 'katma deger' in t and 'karsit inceleme' in t and 'tutana' in t:
                bas.append(i)
            elif 'konu' in t and 'bilgi isteme' in t:
                j = i
                for k in range(i - 1, max(i - 4, -1), -1):
                    if els[k].tag != qn('w:p'):
                        break
                    tk = _ascii_kucuk(Paragraph(els[k], doc).text).strip()
                    if tk.startswith('sayi'):
                        j = k
                        break
                    if tk:                       # 'Sayı' olmayan dolu satıra çarparsak dur
                        break
                bas.append(j)
    if not bas:
        bas = [0]
    else:
        bas = sorted(set(bas))
        bas[0] = 0        # ilk bloğu belge başından başlat (antet/boş satır kaybolmasın)
    bloklar = []
    for k, bi in enumerate(bas):
        son = bas[k + 1] if k + 1 < len(bas) else len(els)
        blok_els = els[bi:son]
        vkn, unvan = None, None
        for el in blok_els:
            if el.tag == qn('w:tbl'):
                tb = Table(el, doc)
                metin = '\n'.join('\t'.join(c.text.strip() for c in r.cells)
                                  for r in tb.rows)
                mk = _ascii_kucuk(metin)
                # Karşı taraf bloğu: tutanakta 'NEZDİNDE…', YMM yazısında
                # 'Hakkında Bilgi İstenilen Mükellef'.
                if 'nezd' in mk or 'hakkinda bilgi istenilen' in mk:
                    vkn, unvan = sablon_vkn_metinden(metin)
                    if vkn:
                        break
        bloklar.append({'ilk': bi, 'son': son, 'vkn': vkn,
                        'unvan': unvan, 'els': blok_els})
    return bloklar

def _docx_blok_belgesi(kaynak_yol, blok_index):
    """Birleşik bir .docx'ten yalnızca `blok_index`. firma bloğunu bırakıp diğer
    blokların gövde elemanlarını silerek tek-firmalık bir docx.Document döndürür
    (stiller/bölüm özellikleri korunur)."""
    import docx
    from docx.oxml.ns import qn
    doc = docx.Document(kaynak_yol)
    bloklar = _docx_firma_bloklari(doc)
    if not bloklar:
        return doc
    if blok_index is None or blok_index >= len(bloklar):
        blok_index = 0
    tut = {id(el) for el in bloklar[blok_index]['els']}
    body = doc.element.body
    for el in list(body):
        if el.tag in (qn('w:p'), qn('w:tbl')) and id(el) not in tut:
            body.remove(el)
    return doc

def _docx_fatura_doldur(doc, firma_df, tum_kolonlar, inceleme_dayanagi=None, log_cb=None):
    """Verilen (tek-firmalık) docx belgesindeki fatura tablosunun veri satırlarını
    firma faturalarıyla doldurur; verilirse İnceleme Dayanağı'nı günceller.
    Yazılan satır sayısını döndürür."""
    from copy import deepcopy
    hedef = None
    for tbl in doc.tables:
        basliklar = []
        for row in tbl.rows[:2]:
            basliklar.extend(c.text.strip() for c in row.cells)
        if _fatura_tablosu_mu(basliklar):
            hedef = tbl
            break
    if hedef is None:
        raise RuntimeError("Şablonda fatura tablosu bulunamadı")

    veri_bas = 0
    for i, row in enumerate(hedef.rows):
        txt = _ascii_kucuk(' '.join(c.text for c in row.cells))
        if txt.strip() and any(k in txt for k in _BAS_KELIME):
            veri_bas = i + 1
        else:
            break
    veri_bas = max(veri_bas, 1)

    # Sütun rollerini başlıklardan çöz (tutanak ve YMM yazısı farklı son sütuna
    # sahip: 'Defter Kayıt' → boş, 'KDV dahil toplam' → matrah+kdv). Roller
    # güvenilmezse satır bazında KONUMSAL eşlemeye düşülür.
    kolon_basliklari = _fatura_kolon_basliklari(hedef.rows, veri_bas)
    roller = [_fatura_kolon_rolu(b) for b in kolon_basliklari]

    proto_tr = None
    if len(hedef.rows) > veri_bas:
        proto_tr = deepcopy(hedef.rows[veri_bas]._tr)
    for row in list(hedef.rows)[veri_bas:]:
        hedef._tbl.remove(row._tr)

    yazilan = 0
    for _, r in firma_df.iterrows():
        hucreler = _fatura_satir_hucreleri(r, tum_kolonlar, roller) \
                   or _word_fatura_satiri(r, tum_kolonlar)
        if proto_tr is not None:
            hedef._tbl.append(deepcopy(proto_tr))
            satir = hedef.rows[-1]
        else:
            satir = hedef.add_row()
        for ci, cell in enumerate(satir.cells):
            if ci < len(hucreler):
                _docx_hucre_yaz(cell, hucreler[ci])
        yazilan += 1

    if inceleme_dayanagi:
        if not _docx_inceleme_dayanagi_yaz(doc, inceleme_dayanagi) and log_cb:
            log_cb("      ⚠️ 'İNCELEME DAYANAĞI' alanı şablonda bulunamadı; sözleşme "
                   "bilgisi güncellenemedi.", "warn")
    return yazilan

def _firma_docx_hazirla(sablon_yol, firma_df, tum_kolonlar, inceleme_dayanagi=None,
                        log_cb=None, blok=None):
    """Şablondan (gerekirse birleşik dosyanın `blok`. bloğunu izole ederek) tek
    firmalık doldurulmuş docx.Document ile yazılan satır sayısını döndürür.
    KAYDETMEZ. python-docx yoksa RuntimeError yükseltir."""
    try:
        import docx  # noqa: F401
    except Exception:
        raise RuntimeError("python-docx yok (pip install python-docx)")
    if blok is None:
        doc = docx.Document(sablon_yol)
    else:
        doc = _docx_blok_belgesi(sablon_yol, blok)
    yazilan = _docx_fatura_doldur(doc, firma_df, tum_kolonlar, inceleme_dayanagi, log_cb)
    return doc, yazilan

def firma_docx_olustur(sablon_yol, firma_df, cikis_yol, tum_kolonlar, log_cb=None,
                       inceleme_dayanagi=None, blok=None):
    """Bir firmanın .docx tutanağını üretip `cikis_yol`'a kaydeder (Word GEREKTİRMEZ).
    `blok` verilirse birleşik şablon dosyasının o firma bloğu kullanılır."""
    doc, yazilan = _firma_docx_hazirla(sablon_yol, firma_df, tum_kolonlar,
                                       inceleme_dayanagi, log_cb, blok=blok)
    os.makedirs(os.path.dirname(os.path.abspath(cikis_yol)), exist_ok=True)
    gercek = _guvenli_docx_kaydet(doc, cikis_yol)
    if log_cb:
        log_cb(f"      📝 Word tutanağı: {Path(gercek).name} ({yazilan} fatura satırı)", "ok")
    return gercek

def _docx_govde_ekle(hedef_doc, kaynak_doc, sayfa_sonu=True):
    """kaynak_doc'un gövdesini (paragraf + tablolar) hedef_doc'un SONUNA, son
    bölüm özelliklerinden (sectPr) ÖNCE ekler. sayfa_sonu=True ise araya sayfa
    sonu koyar. Şablonlar aynı kökten olduğu için stiller uyumludur."""
    from copy import deepcopy
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    body = hedef_doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    def _ekle(el):
        if sectPr is not None:
            sectPr.addprevious(el)
        else:
            body.append(el)
    if sayfa_sonu:
        p = OxmlElement('w:p'); r = OxmlElement('w:r'); br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page'); r.append(br); p.append(r)
        _ekle(p)
    for el in list(kaynak_doc.element.body):
        if el.tag == qn('w:sectPr'):
            continue
        _ekle(deepcopy(el))

def firmalar_tek_docx(bloklar, cikis_yol):
    """Doldurulmuş firma docx.Document'lerini (bloklar) tek bir .docx'te, her firma
    yeni sayfada olacak şekilde birleştirip kaydeder. Kaydedilen yolu döndürür."""
    if not bloklar:
        return None
    hedef = bloklar[0]
    for nd in bloklar[1:]:
        _docx_govde_ekle(hedef, nd, sayfa_sonu=True)
    os.makedirs(os.path.dirname(os.path.abspath(cikis_yol)), exist_ok=True)
    return _guvenli_docx_kaydet(hedef, cikis_yol)

def _sablon_yol_blok(sablon_kaydi):
    """İndeks değerini (yol, blok) olarak çözer. Geriye dönük uyum için düz yol
    (str) da kabul edilir → (yol, None)."""
    if isinstance(sablon_kaydi, (tuple, list)):
        yol = sablon_kaydi[0]
        blok = sablon_kaydi[1] if len(sablon_kaydi) > 1 else None
        return str(yol), blok
    return str(sablon_kaydi), None

def firma_word_uret(sablon_kaydi, firma_df, cikis_kl, sira_no, donem, vkn, temiz,
                    tum_kolonlar, log_cb=None, inceleme_dayanagi=None):
    """Eşleşen şablondan firma tutanağı üretir; uzantıya göre doğru yöntemi seçer:
      .docx → python-docx (Word gerektirmez),  .doc → Word COM (Windows + Word).
    `sablon_kaydi` (yol, blok) olabilir: çok-firmalı .docx'te blok o firmanın
    tutanak sayfasıdır. `inceleme_dayanagi` verilirse sözleşme alanı da güncellenir."""
    yol, blok = _sablon_yol_blok(sablon_kaydi)
    ext = Path(yol).suffix.lower()
    ad = f"{sira_no}) {donem.replace('.','_')}_{vkn}_{temiz}{ext}"
    cikis = str(Path(cikis_kl) / ad)
    if ext == '.docx':
        return firma_docx_olustur(yol, firma_df, cikis, tum_kolonlar, log_cb,
                                  inceleme_dayanagi=inceleme_dayanagi, blok=blok)
    return firma_word_olustur(yol, firma_df, cikis, tum_kolonlar, log_cb,
                              inceleme_dayanagi=inceleme_dayanagi)

def sablon_uretilebilir_mi(sablon_kaydi):
    """Bu şablon türü için üretim yapılabilir mi? (.docx→python-docx, .doc→Word COM)"""
    yol, _ = _sablon_yol_blok(sablon_kaydi)
    if yol.lower().endswith('.docx'):
        return docx_destekli()
    return word_destekli()

# ══════════════════════════════════════════
#  ÖZET RAPOR
# ══════════════════════════════════════════
def ozet_rapor_olustur(df, secilen, df_gecersiz, esik_tek, esik_toplam,
                       yuzde80, donem, basarili, hatali_sayisi, bos_sayisi=0):
    """Bir çalışmanın tek sayfalık özetini (kapsam, hedef, gerçek %, geçersiz
    tutar vb.) hesaplayıp döndürür. Sayılar tutanaklarla aynı kaynaktan
    (para_deger ile) üretilir; iş kurallarını yeniden hesaplamaz, sadece
    zaten seçilmiş sonucu raporlar. Bir openpyxl Workbook döndürür."""
    tutar_col = sutun_bul(list(df.columns), ['kdv hariç tutarı', 'faturanın tutarı'])

    def _topla(frame):
        if frame is None or tutar_col is None or len(frame) == 0:
            return 0.0
        return float(frame[tutar_col].apply(lambda v: para_deger(v) or 0.0).sum())

    toplam_liste  = _topla(df)
    secilen_tutar = sum(_topla(grp) for grp, _neden in secilen.values())
    gecersiz_tutar= _topla(df_gecersiz)
    hedef_tutar   = toplam_liste * (yuzde80 / 100.0)
    kapsam_pct    = (secilen_tutar / toplam_liste * 100.0) if toplam_liste else 0.0

    satirlar = [
        ("Dönem",                         donem),
        ("Tek fatura limiti (₺)",         esik_tek),
        ("Toplam fatura limiti (₺)",      esik_toplam),
        ("Hedef kapsam yüzdesi (%)",      yuzde80),
        ("Liste toplam tutarı (₺)",       round(toplam_liste, 2)),
        (f"Hedef tutar (%{yuzde80:.0f}) (₺)", round(hedef_tutar, 2)),
        ("Seçilen firma sayısı",          len(secilen)),
        ("Tutanağı oluşturulan firma",    basarili),
        ("Oluşturulamayan firma",         hatali_sayisi),
        ("Boş şablonla üretilen firma",   bos_sayisi),
        ("Seçilenlerin toplam tutarı (₺)", round(secilen_tutar, 2)),
        ("GERÇEK KAPSAM (%)",             round(kapsam_pct, 1)),
        ("Geçersiz kimlikli satır sayısı", 0 if df_gecersiz is None else len(df_gecersiz)),
        ("Geçersiz satırların tutarı (₺)", round(gecersiz_tutar, 2)),
    ]

    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Özet"
    kalin = openpyxl.styles.Font(bold=True)
    baslik = ws.cell(1, 1, value="KARŞIT İNCELEME — ÇALIŞMA ÖZETİ"); baslik.font = kalin
    ws.cell(2, 1, value="Ölçüt").font = kalin
    ws.cell(2, 2, value="Değer").font = kalin
    for r, (etiket, deger) in enumerate(satirlar, start=3):
        ws.cell(r, 1, value=etiket)
        h = ws.cell(r, 2, value=deger)
        if isinstance(deger, float):
            h.number_format = '#,##0.00'
    # Gerçek kapsam satırını vurgula (hedefin altındaysa dikkat çeksin)
    kapsam_satir = 3 + [e for e, _ in satirlar].index("GERÇEK KAPSAM (%)")
    ws.cell(kapsam_satir, 1).font = kalin
    ws.cell(kapsam_satir, 2).font = kalin
    ws.column_dimensions['A'].width = 34
    ws.column_dimensions['B'].width = 24
    return wb, kapsam_pct

# ══════════════════════════════════════════
#  ANA İŞLEM
# ══════════════════════════════════════════
def dosyalari_isle(kaynak, esik_tek, esik_toplam, yuzde80, _ekrana_log, tamam_cb,
                   ilerleme_cb=None, cikis_kok=None, pdf_uret=False,
                   sablon_klasor=None, cikti_turu='ikisi', inceleme_dayanagi=None,
                   word_tek_dosya=False, bos_sablon=None):
    # word_tek_dosya: üretilen .docx tutanakları tek bir dosyada (her firma yeni
    #   sayfada) birleştir. bos_sablon: eşleşmeyen firmalar için kullanılacak boş
    #   .docx şablonu (fatura + bilinen ünvan/VKN doldurulur, gerisi kullanıcıda).
    # ilerleme_cb(tamamlanan, toplam): GUI ilerleme çubuğunu günceller.
    # None geçilirse (ör. başsız test) hiçbir şey yapmaz.
    # cikis_kok: çıktı klasörünün üst dizini (None → kaynak dosyanın yanı).
    # pdf_uret: her firma için Excel'in yanına okunur bir PDF kopyası da üret.
    # sablon_klasor: hazır .doc/.docx Word şablonlarının klasörü. Verilirse seçilen
    #   firmalar VKN ile eşleştirilir ve eşleşen şablonun fatura tablosu
    #   güncellenerek Word tutanağı da üretilir.
    # cikti_turu: 'excel' (yalnız Excel), 'word' (yalnız Word — şablon gerekir),
    #   'ikisi' (Excel + eşleşen firmalar için Word). Varsayılan 'ikisi'.
    excel_iste = cikti_turu in ('excel', 'ikisi')
    word_iste  = cikti_turu in ('word', 'ikisi') and bool(sablon_klasor)
    def _ilerle(t, top):
        if ilerleme_cb:
            try: ilerleme_cb(t, top)
            except Exception: pass

    # Tüm günlük satırlarını sakla ki sonunda kalıcı bir .txt olarak yazılabilsin
    # (denetim izi). log_cb burada kasıtlı olarak gölgelenir; içerideki tüm
    # log_cb(...) çağrıları hem ekrana hem bu tampona yazar.
    _gunluk_kayit = []
    def log_cb(msg, tip=''):
        _gunluk_kayit.append(str(msg))
        _ekrana_log(msg, tip)

    def _gunlugu_yaz(klasor, donem):
        try:
            yol = Path(klasor) / f"ISLEM_GUNLUGU_{donem.replace('.','_')}.txt"
            baslik = (f"e-YMM Karşıt İnceleme Asistanı v{SURUM}\n"
                      f"Çalışma zamanı: {datetime.now():%d.%m.%Y %H:%M:%S}\n"
                      f"Kaynak dosya: {Path(kaynak).name}\n"
                      f"Kriterler: tek≥{esik_tek:,.0f}  toplam≥{esik_toplam:,.0f}  "
                      f"kapsam=%{float(yuzde80):.0f}\n" + "="*60 + "\n")
            with open(yol, 'w', encoding='utf-8') as f:
                f.write(baslik + "\n".join(_gunluk_kayit) + "\n")
            return yol
        except Exception:
            return None

    try:
        _ilerle(0, 1)
        log_cb("📂 Dosya okunuyor...", "info")
        df    = ana_listeyi_oku(kaynak)
        donem = donem_bul(Path(kaynak).stem, df)

        # ── İşlem öncesi ÖN BİLGİ (yanlış dosya/kolon baştan yakalansın) ──
        log_cb(f"{'─'*50}", "info")
        log_cb("📋 ÖN BİLGİ (işlemden önce kontrol edin):", "info")
        log_cb(f"   Satır sayısı: {len(df)}   |   Dönem: {donem}", "info")
        for etiket, sut in bulunan_sutunlar(df).items():
            if sut:
                log_cb(f"   ✔ {etiket:10}→ \"{str(sut)[:45]}\"", "ok")
            else:
                log_cb(f"   ✘ {etiket:10}→ (bulunamadı)", "warn")

        # ── Doğruluk uyarıları (yalnızca uyarır, seçimi etkilemez) ──
        for m, t in kdv_tutarlilik_kontrol(df):
            log_cb(m, t)
        mukerrer = mukerrer_fatura_bul(df)
        if mukerrer:
            log_cb(f"  ⚠️  {len(mukerrer)} mükerrer fatura (aynı VKN + fatura no) "
                   f"birden çok satırda görünüyor:", "warn")
            for v, f, adet in mukerrer[:10]:
                log_cb(f"     • VKN {v}  fatura {f}  ×{adet}", "warn")
            if len(mukerrer) > 10:
                log_cb(f"     … ve {len(mukerrer)-10} tane daha", "warn")
        t_toplam, t_disi = donem_disi_tarih_kontrol(df, donem)
        if t_toplam and t_disi / t_toplam > 0.3:
            log_cb(f"  ⚠️  {t_disi}/{t_toplam} faturanın tarihi seçilen dönem "
                   f"({donem}) dışında — yanlış dönem dosyası olabilir.", "warn")
        log_cb(f"{'─'*50}", "info")

        log_cb("🔍 Firmalar filtreleniyor...", "info")

        secilen, df_gecersiz = firmalari_filtrele(df, esik_tek, esik_toplam, yuzde80, log_cb)

        if not secilen:
            log_cb("⚠️  Hiçbir firma kriterleri karşılamıyor.", "warn")
            tamam_cb(None, 0, 0); return

        cikis_taban = Path(cikis_kok) if cikis_kok else Path(kaynak).parent
        cikis_kl = cikis_taban / "Hazır Tutanaklar"
        cikis_kl.mkdir(parents=True, exist_ok=True)

        # Klasörde eski Excel dosyası varsa uyar
        eski_dosyalar = list(cikis_kl.glob("*.xlsx"))
        if eski_dosyalar:
            log_cb(f"⚠️  Klasörde {len(eski_dosyalar)} eski dosya var.", "warn")
            log_cb(f"   Üzerine yazmamak için klasör yeniden adlandırılıyor...", "warn")
            zaman_damgasi = datetime.now().strftime("%Y%m%d_%H%M%S")
            yedek_kl = cikis_kl.parent / f"Hazır Tutanaklar_{zaman_damgasi}"
            cikis_kl.rename(yedek_kl)
            log_cb(f"   Eski klasör: {yedek_kl.name}", "warn")
            cikis_kl = cikis_taban / "Hazır Tutanaklar"
            cikis_kl.mkdir(parents=True, exist_ok=True)
            log_cb(f"   Yeni klasör oluşturuldu.", "ok")

        log_cb(f"📁 Klasör: {cikis_kl}", "info")
        # PDF isteniyor ama reportlab yoksa: kullanıcıyı bir kez uyar, Excel'e devam et
        if pdf_uret and not pdf_destekli():
            log_cb("  ⚠️  PDF üretimi için 'reportlab' kurulu değil; yalnızca Excel "
                   "üretilecek. (pip install reportlab)", "warn")
            pdf_uret = False
        if pdf_uret:
            log_cb("  🧾 PDF kopyalar da üretilecek.", "info")

        # Çıktı türünü bildir
        _mod = {'excel': 'yalnız Excel', 'word': 'yalnız Word',
                'ikisi': 'Excel + (eşleşen firmalar için) Word'}.get(cikti_turu, 'Excel + Word')
        log_cb(f"🧾 Çıktı türü: {_mod}", "info")
        if cikti_turu == 'word' and not sablon_klasor:
            log_cb("  ⚠️  'Yalnız Word' seçildi ama şablon klasörü seçilmedi — "
                   "hiçbir tutanak üretilemeyecek. Lütfen şablon klasörünü seçin.", "err")

        # ── Word şablonları: gerekiyorsa VKN'ye göre indeksle ──
        sablon_index = {}
        if word_iste:
            log_cb(f"🗂  Word şablonları taranıyor: {sablon_klasor}", "info")
            try:
                sablon_index = sablonlari_indeksle(sablon_klasor, log_cb)
            except Exception as e:
                log_cb(f"  ⚠️ Şablon klasörü okunamadı: {e}", "warn")
            log_cb(f"  {len(sablon_index)} şablon VKN ile indekslendi.", "ok")
            if inceleme_dayanagi:
                log_cb(f"  🗓 İnceleme Dayanağı her Word tutanağına yazılacak: "
                       f"\"{inceleme_dayanagi}\"", "info")
            # Şablon uzantılarına göre üretim ön koşulunu kontrol et
            idx_docx = any(_sablon_yol_blok(v)[0].lower().endswith('.docx')
                           for v in sablon_index.values())
            idx_doc  = any(_sablon_yol_blok(v)[0].lower().endswith('.doc')
                           for v in sablon_index.values())
            if idx_docx and not docx_destekli():
                log_cb("  ⚠️  .docx şablonlar için 'python-docx' gerekli; kurulu "
                       "değil (pip install python-docx). Bu şablonlar üretilemez.", "warn")
            if idx_doc and not word_destekli():
                log_cb("  ⚠️  .doc şablonlar için Word otomasyonu (pywin32 + Windows "
                       "Word) gerekli; yok. Bu şablonlar üretilemez. (.docx şablonlar "
                       "Word'süz üretilebilir.)", "warn")
        word_eslesen = []     # (vkn, unvan) — şablonu bulunan firmalar
        word_uretilen = []    # vkn — Word tutanağı gerçekten üretilen firmalar
        word_sablonsuz = []   # (vkn, unvan) — seçilmiş, şablonu yok ve boş şablon da yok
        word_bloklar = []     # (sira, vkn, unvan, doc) — tek dosyada birleştirmek için
        bos_uretilen = []     # (vkn, unvan) — boş/yedek şablonla üretilen firmalar
        if bos_sablon and word_iste:
            log_cb(f"  🆕 Eşleşmeyen firmalar için boş şablon kullanılacak: "
                   f"{Path(bos_sablon).name}", "info")
        if word_tek_dosya and word_iste:
            log_cb("  🧩 Word tutanakları tek dosyada birleştirilecek (yalnızca .docx).", "info")
        log_cb(f"{'─'*50}", "info")

        unvan_col = sutun_bul(list(df.columns), ['satıcının adı', 'ünvanı'])
        fno_col   = sutun_bul(list(df.columns), ['alış faturasının sıra no'])
        basarili  = 0; hatali = []
        vkn_sirali = []   # tutanağı oluşturulan firmaların VKN'leri (dosya sırasıyla)
        sira_no = 0       # yalnızca başarılı dosyalar için ARDIŞIK numara (atlama olmaz)

        toplam_firma = len(secilen)
        _ilerle(0, toplam_firma)
        for islenen, (vkn, (grp, neden)) in enumerate(secilen.items(), start=1):
            uv    = grp[unvan_col].dropna().unique() if unvan_col else []
            unvan = str(uv[0]) if len(uv) else ''
            # Firmadan örnek bir fatura numarası (birden fazla olsa da biri yeterli)
            ornek_fno = ''
            if fno_col:
                fdiz = grp[fno_col].dropna()
                if len(fdiz):
                    ornek_fno = str(fdiz.iloc[0]).strip()
            temiz = dosya_adi_temizle(unvan) if unvan else vkn
            sira_no += 1              # aday numara; hiçbir çıktı üretilmezse geri alınır
            uretildi = False          # bu firma için en az bir dosya üretildi mi?
            hata_mesaji = None        # Excel/Word üretim hatası (varsa)

            # ── Excel tutanağı ──
            if excel_iste:
                try:
                    ad = f"{sira_no}) {donem.replace('.','_')}_{vkn}_{temiz}.xlsx"
                    kayitli_yol = firma_excel_olustur(grp, str(cikis_kl / ad), list(df.columns))
                    uretildi = True
                    if pdf_uret:      # PDF, Excel'in okunur kopyası — hata Excel'i etkilemez
                        try:
                            firma_pdf_olustur(grp, str(Path(kayitli_yol).with_suffix('.pdf')),
                                              list(df.columns), vkn=vkn, unvan=unvan, donem=donem)
                        except Exception as pe:
                            log_cb(f"      ⚠️ PDF üretilemedi ({vkn}): {pe}", "warn")
                except Exception as e:
                    hata_mesaji = str(e)

            # ── Word tutanağı (VKN ile eşleşen şablondan) ──
            if word_iste:
                sy = sablon_index.get(vkn)
                if sy:
                    word_eslesen.append((vkn, unvan))
                    yol_, blok_ = _sablon_yol_blok(sy)
                    if sablon_uretilebilir_mi(sy):
                        try:
                            if word_tek_dosya and yol_.lower().endswith('.docx'):
                                d_, _y = _firma_docx_hazirla(yol_, grp, list(df.columns),
                                            inceleme_dayanagi, log_cb, blok=blok_)
                                word_bloklar.append((sira_no, vkn, unvan, d_))
                            else:
                                firma_word_uret(sy, grp, cikis_kl, sira_no, donem,
                                                vkn, temiz, list(df.columns), log_cb,
                                                inceleme_dayanagi=inceleme_dayanagi)
                            word_uretilen.append(vkn)
                            uretildi = True
                        except Exception as we:
                            log_cb(f"      ⚠️ Word tutanağı üretilemedi ({vkn}): {we}", "warn")
                            if not excel_iste and hata_mesaji is None:
                                hata_mesaji = str(we)
                elif bos_sablon and str(bos_sablon).lower().endswith('.docx') and docx_destekli():
                    # Şablonu yok → boş/yedek şablondan üret (fatura + bilinen ünvan/VKN)
                    try:
                        d_, _y = _firma_docx_hazirla(bos_sablon, grp, list(df.columns),
                                    inceleme_dayanagi, log_cb, blok=None)
                        _docx_nezdinde_yaz(d_, unvan, vkn)
                        if word_tek_dosya:
                            word_bloklar.append((sira_no, vkn, unvan, d_))
                        else:
                            ad = f"{sira_no}) {donem.replace('.','_')}_{vkn}_{temiz}_BOS.docx"
                            _guvenli_docx_kaydet(d_, str(cikis_kl / ad))
                        bos_uretilen.append((vkn, unvan))
                        uretildi = True
                        log_cb(f"      🆕 Boş şablondan tutanak üretildi ({vkn}) — "
                               f"firma bilgilerini fatura üzerinden kontrol edin.", "warn")
                    except Exception as be:
                        log_cb(f"      ⚠️ Boş şablon üretilemedi ({vkn}): {be}", "warn")
                        word_sablonsuz.append((vkn, unvan))
                else:
                    word_sablonsuz.append((vkn, unvan))

            # ── Sonuç: ardışık numaralandırma yalnızca üretilen firmalar için ──
            if uretildi:
                basarili += 1
                vkn_sirali.append((sira_no, vkn, unvan, ornek_fno))
                log_cb(f"  [{sira_no:3}/{len(secilen)}] ✔ {vkn}  {unvan[:35]}", "ok")
                if hata_mesaji:       # kısmi: Word üretildi ama Excel üretilemedi
                    log_cb(f"      ⚠️ Excel üretilemedi ({vkn}): {hata_mesaji}", "warn")
            else:
                sira_no -= 1          # üretilmedi → numarayı geri al (atlama olmasın)
                if hata_mesaji:
                    hatali.append((vkn, unvan, hata_mesaji))
                    log_cb(f"  ✘ HATA  {vkn}  {unvan[:30]}: {hata_mesaji}", "err")
            _ilerle(islenen, toplam_firma)

        log_cb(f"\n{'═'*50}", "info")
        log_cb(f"✅ {basarili}/{len(secilen)} firma tamamlandı."
               + (f"  ⚠️ {len(hatali)} firma oluşturulamadı!" if hatali else ""), "ok")

        # ── Word tutanaklarını tek dosyada birleştir (.docx) ──
        if word_tek_dosya and word_bloklar:
            try:
                birlesik_ad = f"KARSIT_INCELEME_TUTANAKLAR_{donem.replace('.','_')}.docx"
                byol = firmalar_tek_docx([d for _s, _v, _u, d in word_bloklar],
                                         str(cikis_kl / birlesik_ad))
                log_cb(f"🧩 {len(word_bloklar)} Word tutanağı TEK dosyada birleştirildi: "
                       f"{Path(byol).name}", "ok")
            except Exception as e:
                log_cb(f"⚠️ Tek dosyada birleştirilemedi: {e}", "warn")

        # ── Word şablon eşleşme özeti + raporu ──
        if word_iste:
            log_cb(f"🗂  Şablon eşleşmesi: {len(word_eslesen)} firma eşleşti, "
                   f"{len(word_uretilen)} Word tutanağı üretildi"
                   + (f"; {len(bos_uretilen)} boş şablonla" if bos_uretilen else "")
                   + (f"; {len(word_sablonsuz)} firmanın şablonu yok."
                      if word_sablonsuz else "."), "ok")
            if bos_uretilen:
                log_cb(f"   🆕 {len(bos_uretilen)} firma için BOŞ ŞABLON oluşturuldu — "
                       f"firma bilgilerini fatura üzerinden girebilirsiniz.", "warn")
            if word_sablonsuz:
                for v, uv in word_sablonsuz[:15]:
                    log_cb(f"   • Şablon yok: {v:15} {uv[:35]}", "warn")
                if len(word_sablonsuz) > 15:
                    log_cb(f"   … ve {len(word_sablonsuz)-15} firma daha", "warn")
            try:
                wyol = str(cikis_kl / f"WORD_ESLESME_{donem.replace('.','_')}.xlsx")
                wbw = openpyxl.Workbook(); wsw = wbw.active; wsw.title = "Şablon Eşleşme"
                for c, b in enumerate(["VKN/TC", "Ünvan", "Durum"], 1):
                    wsw.cell(1, c, value=b).font = openpyxl.styles.Font(bold=True)
                r = 2
                for v, uv in word_eslesen:
                    wsw.cell(r, 1, value=str(v)).number_format = '@'
                    wsw.cell(r, 2, value=str(uv))
                    wsw.cell(r, 3, value=("Word üretildi" if v in word_uretilen
                                          else "Eşleşti (Word üretilmedi)"))
                    r += 1
                for v, uv in bos_uretilen:
                    wsw.cell(r, 1, value=str(v)).number_format = '@'
                    wsw.cell(r, 2, value=str(uv))
                    wsw.cell(r, 3, value="Boş şablon oluşturuldu (firma bilgisi girilecek)")
                    r += 1
                for v, uv in word_sablonsuz:
                    wsw.cell(r, 1, value=str(v)).number_format = '@'
                    wsw.cell(r, 2, value=str(uv))
                    wsw.cell(r, 3, value="Şablon yok")
                    r += 1
                wsw.column_dimensions['A'].width = 16
                wsw.column_dimensions['B'].width = 45
                wsw.column_dimensions['C'].width = 26
                wbw.save(wyol)
                log_cb(f"🗂  Şablon eşleşme raporu: WORD_ESLESME_{donem.replace('.','_')}.xlsx", "ok")
            except Exception as e:
                log_cb(f"⚠️ Şablon eşleşme raporu yazılamadı: {e}", "warn")

        # ── Oluşturulamayan firmaları belirgin şekilde raporla (sessizce kaybolmasın) ──
        if hatali:
            log_cb(f"{'─'*50}", "err")
            log_cb(f"⚠️  AŞAĞIDAKİ {len(hatali)} FİRMANIN TUTANAĞI OLUŞTURULAMADI:", "err")
            for v, uv, sebep in hatali:
                log_cb(f"   • {v:15} {uv[:35]:35} → {sebep}", "err")
            try:
                import openpyxl as _opx
                hyol = str(cikis_kl / f"OLUSTURULAMAYANLAR_{donem.replace('.','_')}.xlsx")
                wbh = _opx.Workbook(); wsh = wbh.active; wsh.title = "Oluşturulamayanlar"
                for c, b in enumerate(["Vergi/TC No", "Ünvan", "Hata Nedeni"], 1):
                    hc = wsh.cell(1, c, value=b); hc.font = _opx.styles.Font(bold=True)
                for r, (v, uv, sebep) in enumerate(hatali, 2):
                    wsh.cell(r, 1, value=str(v)).number_format = '@'
                    wsh.cell(r, 2, value=str(uv))
                    wsh.cell(r, 3, value=str(sebep))
                wsh.column_dimensions['A'].width = 16
                wsh.column_dimensions['B'].width = 45
                wsh.column_dimensions['C'].width = 60
                wbh.save(hyol)
                log_cb(f"📄 Detay: OLUSTURULAMAYANLAR_{donem.replace('.','_')}.xlsx", "err")
            except Exception:
                pass

        # ── Tutanağı oluşturulan firmaların VKN listesi (dosya sırasıyla, büyükten küçüğe) ──
        if vkn_sirali:
            try:
                vkn_yolu = str(cikis_kl / f"VKN_LISTESI_{donem.replace('.','_')}.xlsx")
                wb_v = openpyxl.Workbook(); ws_v = wb_v.active; ws_v.title = "VKN Listesi"
                basliklar = ["Sıra No", "Vergi Kimlik / TC Kimlik No", "Ünvan", "Örnek Fatura No"]
                for c, bsl in enumerate(basliklar, start=1):
                    hc = ws_v.cell(1, c, value=bsl); hc.font = openpyxl.styles.Font(bold=True)
                for r, (sira, v, uv, fno) in enumerate(vkn_sirali, start=2):
                    ws_v.cell(r, 1, value=sira).number_format = '@'
                    ws_v.cell(r, 2, value=str(v)).number_format = '@'   # baştaki sıfırlar korunur
                    ws_v.cell(r, 3, value=str(uv) if uv else '')
                    ws_v.cell(r, 4, value=str(fno) if fno else '').number_format = '@'
                # Sütun genişlikleri
                ws_v.column_dimensions['A'].width = 8
                ws_v.column_dimensions['B'].width = 26
                ws_v.column_dimensions['C'].width = 45
                ws_v.column_dimensions['D'].width = 22
                wb_v.save(vkn_yolu)
                log_cb(f"📇 VKN listesi: VKN_LISTESI_{donem.replace('.','_')}.xlsx "
                       f"({len(vkn_sirali)} firma, dosya sırasıyla)", "ok")
            except Exception as e:
                log_cb(f"⚠️ VKN listesi yazılamadı: {e}", "warn")

        # ── Geçersiz satır raporu ──
        if df_gecersiz is not None and len(df_gecersiz) > 0:
            try:
                rapor_yolu = str(cikis_kl / f"GECERSIZ_SATIRLAR_{donem.replace('.','_')}.xlsx")
                # Geçersiz satırların hangi VKN/TC içerdiğini de göster
                vkn_col_r = sutun_bul(list(df.columns), ['vergi kimlik','vkn','tc kimlik'])
                unvan_col_r = sutun_bul(list(df.columns), ['satıcının adı','ünvanı'])
                rapor_df = df_gecersiz.copy()
                rapor_df.insert(0, 'Sorun', rapor_df[vkn_col_r].apply(
                    lambda x: _gecersizlik_nedeni(x)) if vkn_col_r else 'Bilinmiyor')
                rapor_df.to_excel(rapor_yolu, index=False)
                log_cb(f"\n📋 Geçersiz satır raporu: GECERSIZ_SATIRLAR_{donem.replace('.','_')}.xlsx", "warn")
                log_cb(f"   {len(df_gecersiz)} satır — VKN/TC kimlik hatalı veya eksik:", "warn")
                # Günlüğe de yaz
                if vkn_col_r:
                    for _, row in df_gecersiz.iterrows():
                        vkn_val = str(row.get(vkn_col_r,'')).strip()
                        unvan_val = str(row.get(unvan_col_r,''))[:40].strip() if unvan_col_r else ''
                        neden = _gecersizlik_nedeni(vkn_val)
                        log_cb(f"   • {vkn_val:15}  {unvan_val:40}  → {neden}", "warn")
            except Exception as e:
                log_cb(f"⚠️  Rapor oluşturulamadı: {e}", "err")

        # ── Çalışma özeti (tek sayfalık kapsam raporu) ──
        try:
            ozet_wb, kapsam_pct = ozet_rapor_olustur(
                df, secilen, df_gecersiz, esik_tek, esik_toplam,
                yuzde80, donem, basarili, len(hatali), bos_sayisi=len(bos_uretilen))
            ozet_yolu = str(cikis_kl / f"OZET_RAPOR_{donem.replace('.','_')}.xlsx")
            ozet_wb.save(ozet_yolu)
            log_cb(f"📊 Özet rapor: OZET_RAPOR_{donem.replace('.','_')}.xlsx "
                   f"(gerçek kapsam %{kapsam_pct:.1f})", "ok")
        except Exception as e:
            log_cb(f"⚠️  Özet rapor yazılamadı: {e}", "warn")

        # ── İşlem günlüğünü kalıcı olarak yaz (denetim izi) ──
        gyol = _gunlugu_yaz(cikis_kl, donem)
        if gyol:
            log_cb(f"📝 İşlem günlüğü: {Path(gyol).name}", "ok")

        log_cb(f"📁 {cikis_kl}", "ok")
        tamam_cb(str(cikis_kl), basarili, len(hatali))

    except Exception as e:
        log_cb(f"\n❌ {e}", "err")
        # Hata olsa bile günlüğü kaydetmeye çalış (kaynağın yanına)
        try:
            _gunlugu_yaz(Path(cikis_kok) if cikis_kok else Path(kaynak).parent, "HATA")
        except Exception:
            pass
        tamam_cb(None, 0, 1)

# ══════════════════════════════════════════
#  GUI
# ══════════════════════════════════════════
APP="#F4F1F2"; KART="#FFFFFF"; KOYU="#241A1C"
KIRMIZI="#A61C2B"; GRI="#7A6E70"; TURUNCU="#C2410C"
ACCENT="#A61C2B"; ACCENT2="#851521"; KENAR="#E7DCDE"; BASLIK_BG="#6E1423"
YESIL="#15803D"; LOG_BG="#1B1417"; SEKME_BG="#EBDEE0"
F_ANA=('Segoe UI',10); F_BLK=('Segoe UI',10,'bold')
F_KUC=('Segoe UI',9);  F_BAS=('Segoe UI',13,'bold')
F_CON=('Consolas',9)

class KDVBolmeApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"YMM Karşıt İnceleme Asistanı  v{SURUM}")
        self.root.geometry("860x650")
        self.root.minsize(820, 600)
        self.root.configure(bg=APP)
        # Pencere başlık çubuğu ikonu (logo.ico varsa)
        try:
            ico = kaynak_yolu('logo.ico')
            if ico.exists():
                self.root.iconbitmap(str(ico))
        except Exception:
            pass
        self._isleniyor = False
        self._logo_img  = None

        # Son kullanılan kriterleri yükle (yoksa varsayılanlar)
        ayar = self._ayar_yukle()
        self.esik_tek    = tk.StringVar(value=ayar.get("esik_tek",    "150000"))
        self.esik_toplam = tk.StringVar(value=ayar.get("esik_toplam", "450000"))
        self.yuzde80     = tk.StringVar(value=ayar.get("yuzde80",     "80"))
        self._cikis_kok  = ayar.get("cikis_kok") or None   # çıktı klasörü (None → kaynağın yanı)
        self.pdf_uret    = tk.BooleanVar(value=bool(ayar.get("pdf_uret", False)))
        self._sablon_klasor = ayar.get("sablon_klasor") or None  # Word şablon klasörü
        # Çıktı türü: 'excel' | 'word' | 'ikisi'
        self.cikti_turu  = tk.StringVar(value=ayar.get("cikti_turu", "excel"))
        # İnceleme Dayanağı (sözleşme) — her yıl değişir; boşsa şablon aynen kalır
        self.inceleme_dayanagi = tk.StringVar(value=ayar.get("inceleme_dayanagi", ""))
        # Word tutanaklarını tek dosyada birleştir + eşleşmeyenler için boş şablon
        self.word_tek_dosya = tk.BooleanVar(value=bool(ayar.get("word_tek_dosya", False)))
        self._bos_sablon = ayar.get("bos_sablon") or None

        self._ui()
        self._surukle_birak()

    @staticmethod
    def _ayar_yukle():
        """Kayıtlı kriterleri okur; dosya yok/bozuksa boş sözlük döner."""
        try:
            with open(AYAR_YOLU, 'r', encoding='utf-8') as f:
                d = json.load(f)
            return d if isinstance(d, dict) else {}
        except Exception:
            return {}

    def _ayar_kaydet(self):
        """Geçerli kriterleri diske yazar; hata olursa sessizce geçilir
        (ayar kaydedememek işlemi engellememeli)."""
        try:
            with open(AYAR_YOLU, 'w', encoding='utf-8') as f:
                json.dump({
                    "esik_tek":    self.esik_tek.get(),
                    "esik_toplam": self.esik_toplam.get(),
                    "yuzde80":     self.yuzde80.get(),
                    "cikis_kok":   self._cikis_kok or "",
                    "pdf_uret":    bool(self.pdf_uret.get()),
                    "sablon_klasor": self._sablon_klasor or "",
                    "cikti_turu":  self.cikti_turu.get(),
                    "inceleme_dayanagi": self.inceleme_dayanagi.get(),
                    "word_tek_dosya": bool(self.word_tek_dosya.get()),
                    "bos_sablon":  self._bos_sablon or "",
                }, f, ensure_ascii=False)
        except Exception:
            pass

    def _stil_kur(self):
        """ttk temasını 'clam'a alıp düz/modern bir görünüm için renkleri ayarlar."""
        try:
            st = ttk.Style()
            try: st.theme_use('clam')
            except Exception: pass
            st.configure("TNotebook", background=APP, borderwidth=0, tabmargins=(2,4,2,0))
            st.configure("TNotebook.Tab", background=SEKME_BG, foreground=GRI,
                         font=('Segoe UI',9,'bold'), padding=(16,8), borderwidth=0)
            st.map("TNotebook.Tab",
                   background=[('selected', APP)],
                   foreground=[('selected', ACCENT)])
            st.configure("TEntry", fieldbackground=KART, background=KART,
                         bordercolor=KENAR, lightcolor=KENAR, darkcolor=KENAR,
                         relief='flat', padding=5)
            st.configure("Ymm.Horizontal.TProgressbar",
                         troughcolor=KENAR, background=ACCENT, borderwidth=0, thickness=8)
            st.configure("Vertical.TScrollbar", background="#8A6E72",
                         troughcolor=LOG_BG, borderwidth=0, arrowcolor="#E7DCDE")
        except Exception:
            pass

    def _buton(self, parent, text, cmd, tur='ikincil'):
        """Düz (flat), hover'lı bir düğme döndürür. tur='ana' → vurgu renkli."""
        if tur == 'ana':
            bg, fg, hov = ACCENT, 'white', ACCENT2
        else:
            bg, fg, hov = KART, KOYU, '#F3E9EA'
        b = tk.Button(parent, text=text, command=cmd, font=('Segoe UI',9),
                      bg=bg, fg=fg, activebackground=hov, activeforeground=fg,
                      relief='flat', bd=0, padx=12, pady=6, cursor='hand2',
                      highlightbackground=KENAR, highlightthickness=1)
        b.bind('<Enter>', lambda e: b.config(bg=hov))
        b.bind('<Leave>', lambda e: b.config(bg=bg))
        return b

    def _ui(self):
        self._stil_kur()

        # ── Üst başlık şeridi ──
        hdr = tk.Frame(self.root, bg=BASLIK_BG, height=62)
        hdr.pack(side='top', fill='x'); hdr.pack_propagate(False)
        mono = tk.Canvas(hdr, width=40, height=40, bg=BASLIK_BG, highlightthickness=0)
        mono.create_oval(2, 2, 38, 38, fill='#F6EFE0', outline='#E3C77A', width=2)
        mono.create_text(20, 21, text="YMM", fill=BASLIK_BG, font=('Segoe UI',9,'bold'))
        mono.pack(side='left', padx=(18,12), pady=11)
        bsol = tk.Frame(hdr, bg=BASLIK_BG); bsol.pack(side='left', pady=(12,0), anchor='w')
        tk.Label(bsol, text="Karşıt İnceleme Asistanı", bg=BASLIK_BG, fg='white',
                 font=('Segoe UI',14,'bold')).pack(anchor='w')
        tk.Label(bsol, text=f"e-YMM  ·  Sürüm {SURUM}", bg=BASLIK_BG, fg='#E4B8BE',
                 font=('Segoe UI',8)).pack(anchor='w')

        # ── Sol panel (ayarlar) ──
        sol = tk.Frame(self.root, bg=APP, width=344)
        sol.pack(side='left', fill='y', padx=(14,8), pady=14)
        sol.pack_propagate(False)
        self.sol = sol

        # ── Ayar sekmeleri (Kriterler / Çıktı / Word Şablon) ──
        nb = ttk.Notebook(sol)
        nb.pack(fill='x', pady=(0,12))

        # — Sekme 1: Kriterler —
        t1 = tk.Frame(nb, bg=APP, padx=10, pady=10); nb.add(t1, text="Kriterler")
        self._esik_satir(t1, "Tek fatura limiti (₺):", self.esik_tek, "ör: 150000")
        self._esik_satir(t1, "Toplam fatura limiti (₺):", self.esik_toplam, "ör: 450000")
        self._esik_satir(t1, "Kapsam yüzdesi (%):", self.yuzde80, "ör: 80")
        tk.Label(t1, text="Tek fatura ≥ limit  VEYA  toplam ≥ limit; yetmezse\n"
                          "büyükten küçüğe ekleyerek % karşılanana dek devam eder.",
                 font=('Segoe UI',8), bg=APP, fg=GRI, justify='left',
                 wraplength=300).pack(fill='x', pady=(8,0))

        # — Sekme 2: Çıktı —
        t2 = tk.Frame(nb, bg=APP, padx=10, pady=10); nb.add(t2, text="Çıktı")
        tk.Label(t2, text="Ne üretilsin?", font=F_KUC, bg=APP, fg=KOYU,
                 anchor='w').pack(fill='x')
        rf = tk.Frame(t2, bg=APP); rf.pack(fill='x', pady=(2,8))
        for etiket, deger in [("Excel", "excel"), ("Word", "word"), ("İkisi", "ikisi")]:
            tk.Radiobutton(rf, text=etiket, value=deger, variable=self.cikti_turu,
                           font=F_KUC, bg=APP, fg=KOYU, activebackground=APP,
                           selectcolor=KART, command=self._ayar_kaydet).pack(side='left', padx=(0,14))
        tk.Checkbutton(t2, text="Excel'in yanına PDF kopya da üret",
                       variable=self.pdf_uret, onvalue=True, offvalue=False,
                       font=F_KUC, bg=APP, fg=KOYU, activebackground=APP,
                       anchor='w', command=self._ayar_kaydet).pack(fill='x', pady=(0,6))
        cf = tk.Frame(t2, bg=APP); cf.pack(fill='x')
        self._buton(cf, "Çıktı klasörü…", self._cikis_klasoru_sec).pack(side='left')
        self.cikis_lbl = tk.Label(cf, text=self._cikis_ozet(), font=('Segoe UI',8),
                                  bg=APP, fg=GRI, anchor='w')
        self.cikis_lbl.pack(side='left', padx=(8,0), fill='x', expand=True)

        # — Sekme 3: Word Şablon —
        t3 = tk.Frame(nb, bg=APP, padx=10, pady=10); nb.add(t3, text="Word Şablon")
        tk.Label(t3, text="Hazır .doc/.docx şablon klasörü (VKN ile eşleşir):",
                 font=F_KUC, bg=APP, fg=KOYU, anchor='w',
                 wraplength=300, justify='left').pack(fill='x')
        wf = tk.Frame(t3, bg=APP); wf.pack(fill='x', pady=(2,8))
        self._buton(wf, "Şablon klasörü…", self._sablon_klasoru_sec).pack(side='left')
        self.sablon_lbl = tk.Label(wf, text=self._sablon_ozet(), font=('Segoe UI',8),
                                   bg=APP, fg=GRI, anchor='w')
        self.sablon_lbl.pack(side='left', padx=(8,0), fill='x', expand=True)
        tk.Label(t3, text="İnceleme Dayanağı (sözleşme) — her Word tutanağına yazılır:",
                 font=F_KUC, bg=APP, fg=KOYU, anchor='w',
                 wraplength=300, justify='left').pack(fill='x')
        ttk.Entry(t3, textvariable=self.inceleme_dayanagi,
                  font=('Segoe UI',9)).pack(fill='x', pady=(2,0))
        tk.Label(t3, text="ör: 31.01.2026 Tarih ve 09 Sayılı Tam Tasdik Sözleşmesi\n"
                          "Boş bırakırsanız şablondaki mevcut yazı aynen kalır.",
                 font=('Segoe UI',8), bg=APP, fg=GRI, justify='left',
                 wraplength=300).pack(fill='x', pady=(3,0))
        try:
            self.inceleme_dayanagi.trace_add('write', lambda *a: self._ayar_kaydet())
        except Exception:
            pass

        # Tek dosyada birleştirme + boş şablon (eşleşmeyenler için)
        tk.Frame(t3, bg=KENAR, height=1).pack(fill='x', pady=8)
        tk.Checkbutton(t3, text="Word tutanaklarını tek dosyada birleştir (.docx)",
                       variable=self.word_tek_dosya, onvalue=True, offvalue=False,
                       font=F_KUC, bg=APP, fg=KOYU, activebackground=APP,
                       anchor='w', command=self._ayar_kaydet,
                       wraplength=300).pack(fill='x')
        bf = tk.Frame(t3, bg=APP); bf.pack(fill='x', pady=(6,0))
        self._buton(bf, "Boş şablon…", self._bos_sablon_sec).pack(side='left')
        self.bos_lbl = tk.Label(bf, text=self._bos_ozet(), font=('Segoe UI',8),
                                bg=APP, fg=GRI, anchor='w')
        self.bos_lbl.pack(side='left', padx=(8,0), fill='x', expand=True)
        tk.Label(t3, text="Eşleşmeyen firma için bu boş şablon kullanılır; fatura ile "
                          "bilinen ünvan/VKN yazılır, kalan bilgileri siz doldurursunuz.",
                 font=('Segoe UI',8), bg=APP, fg=GRI, justify='left',
                 wraplength=300).pack(fill='x', pady=(3,0))

        # ── Sürükle-bırak (davetkâr, vurgu kenarlı kart) ──
        self.birak = tk.Frame(sol, bg=KART, relief='flat', bd=0, cursor='hand2',
                              highlightbackground=ACCENT, highlightcolor=ACCENT,
                              highlightthickness=2)
        self.birak.pack(fill='both', expand=True, pady=(2,10))

        self.birak_ikon = tk.Label(self.birak, text="📥",
                                   font=('Segoe UI',36), bg=KART, fg=ACCENT)
        self.birak_ikon.pack(expand=True, pady=(22,2))

        self.birak_yazi = tk.Label(self.birak,
                                   text="Liste dosyalarını buraya sürükleyin",
                                   font=('Segoe UI',11,'bold'), bg=KART, fg=KOYU,
                                   wraplength=300, justify='center')
        self.birak_yazi.pack(expand=True, pady=(0,2))

        self.birak_alt = tk.Label(self.birak,
                                  text="ya da tıklayarak seçin\nExcel · CSV · TXT  (çoklu seçilebilir)",
                                  font=('Segoe UI',8), bg=KART, fg=GRI,
                                  wraplength=300, justify='center')
        self.birak_alt.pack(pady=(0,16))

        for w in (self.birak, self.birak_ikon, self.birak_yazi, self.birak_alt):
            w.bind('<Button-1>', self._tiklayarak_sec)

        self.durum_lbl = tk.Label(sol, text="● Hazır",
                                  font=('Segoe UI',9), bg=APP, fg=YESIL)
        self.durum_lbl.pack(anchor='w')

        # ── İlerleme çubuğu (işlem sırasında firma sayısına göre dolar) ──
        try:
            self.ilerleme_var = tk.DoubleVar(value=0)
            self.ilerleme_bar = ttk.Progressbar(
                sol, style="Ymm.Horizontal.TProgressbar",
                orient='horizontal', mode='determinate',
                variable=self.ilerleme_var, maximum=100)
            self.ilerleme_bar.pack(fill='x', pady=(6, 0))
            self.ilerleme_yazi = tk.Label(sol, text="", font=('Segoe UI', 8),
                                          bg=APP, fg=GRI)
            self.ilerleme_yazi.pack(anchor='w')
        except Exception:
            self.ilerleme_var = None
            self.ilerleme_bar = None
            self.ilerleme_yazi = None

        # ── Sağ panel (log) ──
        sag = tk.Frame(self.root, bg=APP)
        sag.pack(side='right', fill='both', expand=True, padx=(0,16), pady=14)

        basr = tk.Frame(sag, bg=APP); basr.pack(fill='x')
        tk.Label(basr, text="İşlem Günlüğü", font=('Segoe UI',10,'bold'),
                 bg=APP, fg=KOYU, anchor='w').pack(side='left')
        tk.Label(basr, text="canlı", font=('Segoe UI',8), bg=APP, fg=GRI).pack(side='right')

        lf = tk.Frame(sag, bg=LOG_BG, highlightbackground=KENAR, highlightthickness=1)
        lf.pack(fill='both', expand=True, pady=(6,0))

        self.log = tk.Text(lf, font=F_CON, bg=LOG_BG, fg='#CBD5E1',
                           relief='flat', bd=10, wrap='word', state='disabled',
                           insertbackground='#CBD5E1', selectbackground='#334155')
        sb = ttk.Scrollbar(lf, command=self.log.yview)
        self.log.configure(yscrollcommand=sb.set)
        sb.pack(side='right', fill='y')
        self.log.pack(fill='both', expand=True)

        self.log.tag_config('ok',   foreground='#4ADE80')
        self.log.tag_config('err',  foreground='#F87171')
        self.log.tag_config('info', foreground='#7DD3FC')
        self.log.tag_config('warn', foreground='#FBBF24')

        self._log("e-YMM Karşıt İnceleme Asistanı hazır.\n"
                  "Soldaki sekmelerden kriter ve çıktı türünü seçin,\n"
                  "sonra listeyi bırakma alanına sürükleyin.", "info")

    def _esik_satir(self, parent, etiket, var, ipucu):
        tk.Label(parent, text=etiket, font=F_KUC, bg=APP,
                 fg=KOYU, anchor='w').pack(fill='x')
        f = tk.Frame(parent, bg=APP); f.pack(fill='x', pady=(2,0))
        ttk.Entry(f, textvariable=var, font=F_ANA, width=14).pack(side='left', padx=(0,6))
        tk.Label(f, text=ipucu, font=('Segoe UI',8),
                 bg=APP, fg='#9CA3AF').pack(side='left')

    def _logo_yukle(self):
        for yol in [kaynak_yolu('logo.png'),
                    kaynak_yolu('logo.jpg'),
                    kaynak_yolu('logo.ico')]:
            if yol.exists():
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(str(yol)).convert('RGBA')
                    img.thumbnail((100,100), Image.LANCZOS)
                    bg = Image.new('RGBA', img.size, (245,246,250,255))
                    bg.paste(img, mask=img.split()[3])
                    self._logo_img = ImageTk.PhotoImage(bg)
                    tk.Label(self.sol, image=self._logo_img, bg=APP).pack(pady=(0,4))
                    return
                except: pass
        c = tk.Canvas(self.sol, width=90, height=90, bg=APP, highlightthickness=0)
        c.pack(pady=(0,4))
        c.create_oval(5,5,85,85, fill=KIRMIZI, outline='')
        c.create_text(45,45, text='YMM', fill='white', font=('Segoe UI',16,'bold'))

    def _surukle_birak(self):
        try:
            import tkinterdnd2
            self.root.drop_target_register(tkinterdnd2.DND_FILES)
            self.root.dnd_bind('<<Drop>>',
                               lambda e: self._isle_coklu(self._dnd_ayikla(e.data)))
        except: pass

    @staticmethod
    def _dnd_ayikla(data):
        """Sürükle-bırak verisinden dosya yollarını ayıklar. Boşluklu yollar
        {..} içinde gelir ve birden çok dosya bırakılabilir."""
        parcalar = re.findall(r'\{([^}]*)\}|(\S+)', str(data))
        return [a or b for a, b in parcalar if (a or b)]

    def _tiklayarak_sec(self, e=None):
        """Sürükle-bırak alanına tıklanınca dosya seçtirir (çoklu seçim)."""
        dosyalar = filedialog.askopenfilenames(
            title="KDV Liste(ler)ini Seçin",
            filetypes=[("Desteklenen", "*.xls *.xlsx *.csv *.txt"),
                       ("Excel", "*.xls *.xlsx"),
                       ("CSV/Metin", "*.csv *.txt"),
                       ("Tümü", "*.*")])
        if dosyalar:
            self._isle_coklu(list(dosyalar))

    def _cikis_ozet(self):
        if self._cikis_kok:
            return f"→ {Path(self._cikis_kok).name or self._cikis_kok}"
        return "(kaynağın yanı)"

    def _cikis_klasoru_sec(self):
        d = filedialog.askdirectory(
            title="Çıktı klasörünü seçin (İptal → kaynak dosyanın yanı)")
        self._cikis_kok = d or None
        try:
            self.cikis_lbl.config(text=self._cikis_ozet())
        except Exception:
            pass
        self._ayar_kaydet()

    def _sablon_ozet(self):
        if self._sablon_klasor:
            return f"→ {Path(self._sablon_klasor).name or self._sablon_klasor}"
        return "(seçilmedi)"

    def _bos_ozet(self):
        if self._bos_sablon:
            return f"→ {Path(self._bos_sablon).name}"
        return "(seçilmedi)"

    def _bos_sablon_sec(self):
        d = filedialog.askopenfilename(
            title="Eşleşmeyen firmalar için boş .docx şablonu seçin",
            filetypes=[("Word .docx", "*.docx"), ("Tümü", "*.*")])
        self._bos_sablon = d or None
        try:
            self.bos_lbl.config(text=self._bos_ozet())
        except Exception:
            pass
        self._ayar_kaydet()

    def _sablon_klasoru_sec(self):
        d = filedialog.askdirectory(
            title="Hazır Word (.doc/.docx) şablonlarının bulunduğu klasörü seçin")
        self._sablon_klasor = d or None
        try:
            self.sablon_lbl.config(text=self._sablon_ozet())
        except Exception:
            pass
        # Klasör seçildi ama çıktı 'yalnız Excel' ise, Word'ü de üretsin diye 'İkisi'ye al
        if self._sablon_klasor and self.cikti_turu.get() == 'excel':
            self.cikti_turu.set('ikisi')
        self._ayar_kaydet()

    def _kriter_al(self):
        """Kriter alanlarını okuyup doğrular. Geçerliyse (True, (et,eto,y)),
        değilse kullanıcıya hata gösterip (False, None) döner."""
        try:
            et  = float(self.esik_tek.get().replace('.','').replace(',','.'))
            eto = float(self.esik_toplam.get().replace('.','').replace(',','.'))
            y   = float(self.yuzde80.get().replace(',','.'))
        except Exception:
            messagebox.showerror("Hata", "Kriter değerleri geçersiz.\nSadece sayı girin.")
            return (False, None)
        ok, mesaj = kriter_dogrula(et, eto, y)
        if not ok:
            messagebox.showerror("Geçersiz kriter", mesaj)
            return (False, None)
        return (True, (et, eto, y))

    def _isle(self, dosya):
        """Tek dosya — toplu işleyiciye yönlendirir (DND/uyum için korunur)."""
        self._isle_coklu([dosya])

    def _isle_coklu(self, dosyalar):
        if self._isleniyor:
            self._log("⚠️  İşlem devam ediyor...", "warn"); return
        FORMATLAR = ('.xls', '.xlsx', '.csv', '.txt')
        gecerli = [d for d in dosyalar if Path(d).suffix.lower() in FORMATLAR]
        for d in dosyalar:
            if Path(d).suffix.lower() not in FORMATLAR:
                self._log(f"❌ Desteklenmeyen format atlandı: {Path(d).name}", "err")
        if not gecerli:
            self._log("❌ İşlenecek geçerli dosya yok (.xls/.xlsx/.csv/.txt).", "err")
            return

        ok, kriter = self._kriter_al()
        if not ok:
            return
        esik_tek, esik_toplam, yuzde80 = kriter
        pdf_uret = bool(self.pdf_uret.get())   # ana thread'de oku, worker'a geçir
        cikti_turu = self.cikti_turu.get()
        inceleme_dayanagi = self.inceleme_dayanagi.get().strip()
        word_tek_dosya = bool(self.word_tek_dosya.get())
        bos_sablon = self._bos_sablon
        # Word gereken modda şablon klasörü şart; yoksa kullanıcıyı uyar
        sablon_klasor = self._sablon_klasor if cikti_turu in ('word', 'ikisi') else None
        if cikti_turu in ('word', 'ikisi') and not sablon_klasor:
            if cikti_turu == 'word':
                messagebox.showerror(
                    "Şablon klasörü gerekli",
                    "'Yalnız Word' seçtiniz ama şablon klasörü seçmediniz.\n"
                    "Lütfen 'Şablon klasörü…' ile hazır .doc/.docx şablonların "
                    "bulunduğu klasörü seçin.")
                return
            self._log("⚠️ 'İkisi' seçili ama şablon klasörü yok; yalnızca Excel "
                      "üretilecek.", "warn")
        self._ayar_kaydet()

        self._isleniyor = True
        self._birak_guncelle("⏳ İşleniyor...", TURUNCU)
        self._ilerleme(0, 0)
        self.durum_lbl.config(
            text=(f"{len(gecerli)} dosya işleniyor…" if len(gecerli) > 1
                  else Path(gecerli[0]).name), fg=KOYU)

        threading.Thread(
            target=self._batch_worker,
            args=(gecerli, esik_tek, esik_toplam, yuzde80, pdf_uret, sablon_klasor,
                  cikti_turu, inceleme_dayanagi, word_tek_dosya, bos_sablon),
            daemon=True
        ).start()

    def _batch_worker(self, dosyalar, esik_tek, esik_toplam, yuzde80, pdf_uret,
                      sablon_klasor=None, cikti_turu='ikisi', inceleme_dayanagi=None,
                      word_tek_dosya=False, bos_sablon=None):
        toplam_b = 0; toplam_h = 0; son_klasor = None
        n = len(dosyalar)
        for i, dosya in enumerate(dosyalar, 1):
            self._log(f"\n{'═'*50}", "info")
            onek = f"[{i}/{n}] " if n > 1 else ""
            self._log(f"📂 {onek}{Path(dosya).name}", "info")
            sonuc = {}
            def _tamam_ic(kl, b, h, _s=sonuc):
                _s['klasor'] = kl; _s['b'] = b; _s['h'] = h
            try:
                dosyalari_isle(dosya, esik_tek, esik_toplam, yuzde80,
                               self._log, _tamam_ic, self._ilerleme,
                               self._cikis_kok, pdf_uret, sablon_klasor, cikti_turu,
                               inceleme_dayanagi or None,
                               word_tek_dosya=word_tek_dosya, bos_sablon=bos_sablon)
            except Exception as e:
                self._log(f"❌ {e}", "err")
            toplam_b += sonuc.get('b', 0); toplam_h += sonuc.get('h', 0)
            if sonuc.get('klasor'):
                son_klasor = sonuc['klasor']
        if n > 1:
            self._log(f"\n{'═'*50}", "info")
            self._log(f"🏁 Toplu işlem bitti: {n} dosya → {toplam_b} tutanak"
                      + (f", {toplam_h} hata" if toplam_h else ""), "ok")
        self._tamam(son_klasor, toplam_b, toplam_h)

    def _log(self, msg, tip=''):
        def _():
            self.log.config(state='normal')
            self.log.insert('end', msg+'\n', tip or '')
            self.log.see('end')
            self.log.config(state='disabled')
        self.root.after(0, _)

    def _ilerleme(self, tamamlanan, toplam):
        """İş parçacığından çağrılır; ilerleme çubuğunu ana thread'de günceller."""
        if self.ilerleme_bar is None:
            return
        yuzde = (tamamlanan / toplam * 100.0) if toplam else 0.0
        def _():
            try:
                self.ilerleme_var.set(yuzde)
                if toplam and tamamlanan < toplam:
                    self.ilerleme_yazi.config(text=f"{tamamlanan}/{toplam} firma")
                elif toplam:
                    self.ilerleme_yazi.config(text=f"{toplam}/{toplam} firma ✓")
                else:
                    self.ilerleme_yazi.config(text="")
            except Exception:
                pass
        self.root.after(0, _)

    def _tamam(self, klasor, basarili, hatali):
        self._isleniyor = False
        def _():
            if klasor and basarili > 0:
                self._birak_guncelle(f"✅ {basarili} dosya hazır", "#68D391")
                self.durum_lbl.config(text=f"{basarili} Excel oluşturuldu", fg='#276749')
                try: os.startfile(klasor)
                except: pass
            else:
                self._birak_guncelle("❌ Hata oluştu", "#FC8181")
                self.durum_lbl.config(text="Hata — log'u inceleyin", fg='#C53030')
            def _sifirla():
                self._birak_guncelle("Yeni dosya için buraya bırakın", GRI)
                if self.ilerleme_bar is not None:
                    try:
                        self.ilerleme_var.set(0)
                        self.ilerleme_yazi.config(text="")
                    except Exception:
                        pass
            self.root.after(4000, _sifirla)
        self.root.after(0, _)

    def _birak_guncelle(self, yazi, renk):
        self.birak_yazi.config(text=yazi, fg=renk)

# ══════════════════════════════════════════
if __name__ == '__main__':
    try:
        import tkinterdnd2; root = tkinterdnd2.Tk()
    except: root = tk.Tk()
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except: pass
    KDVBolmeApp(root)
    root.mainloop()
