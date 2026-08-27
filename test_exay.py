"""exay.py için otomatik test paketi.

CLAUDE.md §8 uyarınca headless çalışır (conftest.py tkinter'ı stub'lar) ve
§2'deki iş kurallarını (2 aşamalı %80 seçimi, %80 paydası = tüm liste, geçersiz
VKN'lerin geçersizliği, ardışık numaralandırma, sayısal tutar/KDV) doğrular.

Gerçek GİB dosyaları bu depoda bulunmadığından, üç liste tipi (yeni GİB, eski
GİB, muhasebe/191) sentetik olarak üretilir ve mantık bunlar üzerinde test edilir.

Çalıştırma:  pytest -q
"""
import openpyxl
import pandas as pd
import pytest

import exay


# ── Sessiz log geri çağrısı ──────────────────────────────────────────────────
def _sessiz(*a, **k):
    pass


# ══════════════════════════════════════════════════════════════════════════
#  Saf yardımcı fonksiyonlar
# ══════════════════════════════════════════════════════════════════════════
@pytest.mark.parametrize("girdi,beklenen", [
    ("1.234.567,89", 1234567.89),   # TR binlik + ondalık
    ("1,234,567.89", 1234567.89),   # EN binlik + ondalık
    ("45927,50", 45927.5),          # TR ondalık
    ("1234", 1234.0),
    (1000, 1000.0),
    (1000.5, 1000.5),
    ("", None),
    (None, None),
    ("abc", None),
    ("1500 ₺", 1500.0),            # para birimi işareti temizlenir
    ("1.234,50 TL", 1234.5),       # TL eki + TR biçim
])
def test_para_deger(girdi, beklenen):
    assert exay.para_deger(girdi) == beklenen


def test_tarih_fmt():
    assert exay.tarih_fmt("2026-04-15") == "15.04.2026"
    assert exay.tarih_fmt("") == ""
    # Tanınmayan biçim olduğu gibi döner
    assert exay.tarih_fmt("15/04/2026") == "15/04/2026"


def test_sayi_fmt():
    assert exay.sayi_fmt(1000) == "1000"      # tam sayı → küsürat gösterilmez
    assert exay.sayi_fmt("1234,5") == "1234,50"  # ondalıklı → 2 hane, virgüllü
    assert exay.sayi_fmt("") == ""


def test_kdv_sutunu_bul():
    # Yeni tip (kesme işaretsiz) doğru seçilir, matrah/toplam KDV'siyle karışmaz
    kols = ["Alış Faturasının KDV Hariç Tutarı", "KDV si", "Toplam İndirilecek KDV"]
    assert exay.kdv_sutunu_bul(kols) == "KDV si"
    # Eski tip (kesme işaretli)
    assert exay.kdv_sutunu_bul(["Matrah", "KDV'si"]) == "KDV'si"
    # Yalnızca yasaklı sütunlar varsa None
    assert exay.kdv_sutunu_bul(["KDV Hariç Tutarı", "Toplam KDV"]) is None


def test_seri_sutunu_bul():
    kols = ["Tarih", "Seri", "No"]
    assert exay.seri_sutunu_bul(kols, "Tarih", "No") == "Seri"
    # Ayrı seri sütunu yok; tarihin sağındaki numara sütunudur → seri yok
    kols2 = ["Tarih", "No"]
    assert exay.seri_sutunu_bul(kols2, "Tarih", "No") is None


def test_donem_bul_turkce_ay():
    # Türkçe karakterli ve ASCII (diakritiksiz) ay adlarının ikisi de tanınmalı
    assert exay.donem_bul("NİSAN_2026_liste") == "04.2026"
    assert exay.donem_bul("NISAN_2026_liste") == "04.2026"   # ASCII yazım
    assert exay.donem_bul("AGUSTOS 2025") == "08.2025"
    # Sayısal ay + yıl
    assert exay.donem_bul("04_2026_kdv") == "04.2026"


def test_gecersizlik_nedeni():
    assert "Boş" in exay._gecersizlik_nedeni("")
    assert "kısa" in exay._gecersizlik_nedeni("123")
    assert "uzun" in exay._gecersizlik_nedeni("123456789012")
    assert "Sayısal değil" in exay._gecersizlik_nedeni("ABC123")
    assert "tutucu" in exay._gecersizlik_nedeni("0000000000")


# ══════════════════════════════════════════════════════════════════════════
#  Sentetik liste üreticileri
# ══════════════════════════════════════════════════════════════════════════
def _yeni_gib_yaz(yol):
    """Yeni GİB tipi liste: başlık 0. satırda, gerçek seri sütunu, 'KDV si'."""
    satirlar = [
        # (tarih, seri, no, matrah, kdv, unvan, vkn)
        ("2026-04-01", "A", "BBK1", 200000, 36000, "FIRMA A", "1000000001"),  # tek≥150k
        ("2026-04-02", "A", "BBK2", 300000, 54000, "FIRMA B", "1000000002"),  # B topl 600k
        ("2026-04-03", "A", "BBK3", 300000, 54000, "FIRMA B", "1000000002"),
        ("2026-04-04", "A", "BBK4", 100000, 18000, "FIRMA C", "1000000003"),  # eşik altı
        ("2026-04-05", "A", "BBK5",  50000,  9000, "FIRMA D", "1000000004"),  # eşik altı
        ("2026-04-06", "A", "BBK6",  40000,  7200, "FIRMA E", "0000000000"),  # geçersiz VKN
    ]
    kols = ["Alış Faturasının Tarihi", "Alış Faturasının Serisi",
            "Alış Faturasının Sıra No'su", "Alış Faturasının KDV Hariç Tutarı",
            "KDV si", "Satıcının Adı-Soyadı / Ünvanı",
            "Satıcının Vergi Kimlik Numarası"]
    df = pd.DataFrame(satirlar, columns=kols)
    df.to_excel(yol, index=False)


def _muhasebe_yaz(yol):
    """Muhasebe (191) dökümü: Borç=KDV, Matrah=matrah, Satıcının sütunu yok."""
    satirlar = [
        ("191.01", "2026-01-10", "F1", "1000000001", "FIRMA A", 36000, 200000),
        ("191.01", "2026-01-11", "F2", "1000000002", "FIRMA B",  9000,  50000),
    ]
    kols = ["Hesap Kodu", "Tarih", "Fatura No", "Vergi Kimlik No",
            "Açıklama", "Borç", "Matrah"]
    df = pd.DataFrame(satirlar, columns=kols)
    df.to_excel(yol, index=False)


# ══════════════════════════════════════════════════════════════════════════
#  İş kuralı testleri — firmalari_filtrele (KALP)
# ══════════════════════════════════════════════════════════════════════════
def test_filtrele_kapsam_ve_secim(tmp_path):
    yol = tmp_path / "nisan.xlsx"
    _yeni_gib_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))

    secilen, gecersiz = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)

    # A (tek fatura) ve B (toplam) seçilmeli; C ve D seçilmemeli
    assert "1000000001" in secilen           # FIRMA A
    assert "1000000002" in secilen           # FIRMA B
    assert "1000000003" not in secilen        # FIRMA C
    assert "1000000004" not in secilen        # FIRMA D

    # Geçersiz VKN'li satır (0000000000) geçersizler listesinde olmalı
    assert len(gecersiz) == 1

    # §2/§3: %80 paydası = TÜM liste (geçersiz tutar dahil).
    # Beklenen gerçek kapsam = (200000+600000) / 990000 = %80.8
    tutar_col = exay.sutun_bul(list(df.columns),
                               ['kdv hariç tutarı', 'faturanın tutarı'])
    toplam_liste = df[tutar_col].apply(lambda v: exay.para_deger(v) or 0).sum()
    secilen_tutar = sum(
        grp[tutar_col].apply(lambda v: exay.para_deger(v) or 0).sum()
        for grp, _ in secilen.values())
    kapsam = secilen_tutar / toplam_liste * 100
    assert toplam_liste == pytest.approx(990000)
    assert kapsam == pytest.approx(80.8, abs=0.1)


def test_filtrele_buyukten_kucuge_sirali(tmp_path):
    yol = tmp_path / "nisan.xlsx"
    _yeni_gib_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))
    secilen, _ = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    # Dosya isimlendirme için büyükten küçüğe sıralı olmalı → B(600k) önce, A(200k) sonra
    anahtarlar = list(secilen.keys())
    assert anahtarlar[0] == "1000000002"   # FIRMA B en büyük
    assert anahtarlar[1] == "1000000001"   # FIRMA A


def test_filtrele_vkn_onde_sifir_tamamlama(tmp_path):
    """8-9 haneli VKN'lerde önde eksik sıfır otomatik tamamlanmalı ve geçerli sayılmalı."""
    satirlar = [
        ("2026-04-01", "A", "N1", 500000, 90000, "KISA VKN FIRMA", "71419747"),  # 8 hane
    ]
    kols = ["Alış Faturasının Tarihi", "Alış Faturasının Serisi",
            "Alış Faturasının Sıra No'su", "Alış Faturasının KDV Hariç Tutarı",
            "KDV si", "Satıcının Adı-Soyadı / Ünvanı",
            "Satıcının Vergi Kimlik Numarası"]
    yol = tmp_path / "kisa.xlsx"
    pd.DataFrame(satirlar, columns=kols).to_excel(yol, index=False)
    df = exay.ana_listeyi_oku(str(yol))
    secilen, gecersiz = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    assert "0071419747" in secilen     # önde sıfır tamamlandı
    assert len(gecersiz) == 0


def test_filtrele_bos_liste_cokmez():
    """Tümü sıfır/boş tutarlı liste bölme hatası vermeden çalışmalı (robustluk)."""
    df = pd.DataFrame({
        "Satıcının Vergi Kimlik Numarası": ["1000000001"],
        "Alış Faturasının KDV Hariç Tutarı": [0],
        "Satıcının Adı-Soyadı / Ünvanı": ["SIFIR FIRMA"],
    })
    secilen, gecersiz = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    assert isinstance(secilen, dict)   # çökmeden döndü


# ══════════════════════════════════════════════════════════════════════════
#  Muhasebe tipi tanıma
# ══════════════════════════════════════════════════════════════════════════
def test_muhasebe_tipi_eslenir(tmp_path):
    yol = tmp_path / "191.xlsx"
    _muhasebe_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))
    # Standart GİB adlarına çevrilmiş olmalı
    assert exay.sutun_bul(list(df.columns), ['kdv hariç tutarı']) is not None
    assert exay.kdv_sutunu_bul(list(df.columns)) is not None
    secilen, _ = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    assert "1000000001" in secilen     # tek fatura 200000 ≥ 150000


# ══════════════════════════════════════════════════════════════════════════
#  Tutanak Excel çıktısı — şablon ve sayısal biçim
# ══════════════════════════════════════════════════════════════════════════
def test_firma_excel_sablon_ve_sayisal(tmp_path):
    yol = tmp_path / "nisan.xlsx"
    _yeni_gib_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))
    secilen, _ = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    grp = secilen["1000000001"][0]     # FIRMA A
    cikti = tmp_path / "firma_a.xlsx"
    exay.firma_excel_olustur(grp, str(cikti), list(df.columns))

    wb = openpyxl.load_workbook(cikti)
    ws = wb.active
    basliklar = [ws.cell(1, c).value for c in range(1, len(exay.SABLON_SUTUNLAR) + 1)]
    assert basliklar == exay.SABLON_SUTUNLAR
    # 4. sütun tutar, 5. sütun KDV → GERÇEK SAYI olmalı (metin değil)
    assert isinstance(ws.cell(2, 4).value, (int, float))
    assert isinstance(ws.cell(2, 5).value, (int, float))
    assert ws.cell(2, 4).value == pytest.approx(200000)


# ══════════════════════════════════════════════════════════════════════════
#  Özet rapor (yeni özellik)
# ══════════════════════════════════════════════════════════════════════════
def test_ozet_rapor(tmp_path):
    yol = tmp_path / "nisan.xlsx"
    _yeni_gib_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))
    secilen, gecersiz = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    wb, kapsam = exay.ozet_rapor_olustur(
        df, secilen, gecersiz, 150000, 450000, 80, "04.2026",
        basarili=len(secilen), hatali_sayisi=0)
    assert kapsam == pytest.approx(80.8, abs=0.1)
    ws = wb.active
    metin = "\n".join(str(ws.cell(r, 1).value) for r in range(1, ws.max_row + 1))
    assert "GERÇEK KAPSAM" in metin


# ══════════════════════════════════════════════════════════════════════════
#  Uçtan uca — dosyalari_isle tutanakları + yan dosyaları üretir, ilerleme çağrılır
# ══════════════════════════════════════════════════════════════════════════
def test_dosyalari_isle_uctan_uca(tmp_path):
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)

    ilerleme_kayit = []
    sonuc = {}

    def _tamam(klasor, basarili, hatali):
        sonuc['klasor'] = klasor
        sonuc['basarili'] = basarili
        sonuc['hatali'] = hatali

    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, _tamam,
                        ilerleme_cb=lambda t, top: ilerleme_kayit.append((t, top)))

    assert sonuc['basarili'] == 2         # A ve B
    assert sonuc['hatali'] == 0
    cikis = tmp_path / "Hazır Tutanaklar"
    dosyalar = [p.name for p in cikis.glob("*.xlsx")]
    # 2 firma tutanağı + yan dosyalar (VKN listesi, geçersiz satırlar, özet)
    assert any(n.startswith("1)") for n in dosyalar)
    assert any(n.startswith("2)") for n in dosyalar)
    assert any(n.startswith("VKN_LISTESI") for n in dosyalar)
    assert any(n.startswith("GECERSIZ_SATIRLAR") for n in dosyalar)
    assert any(n.startswith("OZET_RAPOR") for n in dosyalar)
    # Kalıcı işlem günlüğü (.txt) yazılmış olmalı (denetim izi)
    assert any(p.name.startswith("ISLEM_GUNLUGU") for p in cikis.glob("*.txt"))
    # İlerleme geri çağrısı en az bir kez çağrılmış ve sona ulaşmış olmalı
    assert ilerleme_kayit and ilerleme_kayit[-1] == (2, 2)


# ══════════════════════════════════════════════════════════════════════════
#  Eski GİB tipi (başlık alt satırda, adsız seri sütunu, kesme işaretli KDV'si)
# ══════════════════════════════════════════════════════════════════════════
def _eski_gib_yaz(yol):
    """Eski GİB tipi: 2 başlık/altbilgi satırı, sonra gerçek başlık; seri sütunu
    adsız (Unnamed), 'KDV'si' kesme işaretli."""
    wb = openpyxl.Workbook(); ws = wb.active
    ws.append(["İNDİRİLECEK KDV LİSTESİ"])
    ws.append(["Dönem: 01/2026"])
    ws.append(["Alış Faturasının Tarihi", "", "Alış Faturasının Sıra No'su",
               "Satıcının Adı-Soyadı / Ünvanı", "Satıcının Vergi Kimlik Numarası",
               "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı", "KDV'si"])
    ws.append(["2026-01-05", "", "F1", "FIRMA X", "1000000010", 500000, 90000])
    ws.append(["2026-01-06", "", "F2", "FIRMA Y", "1000000011", 50000, 9000])
    wb.save(yol)


def test_eski_gib_okuma_ve_secim(tmp_path):
    yol = tmp_path / "OCAK_2026.xlsx"
    _eski_gib_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))
    # Başlık alt satırdan doğru bulunmalı; kritik sütunlar eşleşmeli
    assert exay.sutun_bul(list(df.columns), ['vergi kimlik']) is not None
    assert exay.kdv_sutunu_bul(list(df.columns)) == "KDV'si"
    assert exay.sutun_bul(list(df.columns), ['kdv hariç tutarı']) is not None
    secilen, _ = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    assert "1000000010" in secilen        # 500.000 ≥ 150.000 (tek fatura)
    assert "1000000011" not in secilen     # 50.000 eşik altı


# ══════════════════════════════════════════════════════════════════════════
#  CSV girdi (Türkçe kodlama + noktalı virgül ayraç + muhasebe tipi)
# ══════════════════════════════════════════════════════════════════════════
def test_csv_muhasebe_okuma(tmp_path):
    yol = tmp_path / "OCAK_2026.csv"
    icerik = (
        "Hesap Kodu;Tarih;Fatura No;Vergi Kimlik No;Açıklama;Borç;Matrah\n"
        "191.01;2026-01-10;F1;0071419747;FIRMA A;36000,00;200000,00\n"
        "191.01;2026-01-11;F2;1000000002;FIRMA B;9000,00;50000,00\n"
    )
    yol.write_text(icerik, encoding="cp1254")   # TR Windows kodlaması
    df = exay.ana_listeyi_oku(str(yol))
    # Muhasebe eşlemesi uygulanmış olmalı
    assert exay.kdv_sutunu_bul(list(df.columns)) is not None
    secilen, _ = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    # 8 haneli VKN önde sıfır tamamlanarak geçerli sayılmalı ve seçilmeli
    assert "0071419747" in secilen


# ══════════════════════════════════════════════════════════════════════════
#  Kriter doğrulama
# ══════════════════════════════════════════════════════════════════════════
@pytest.mark.parametrize("et,eto,y,ok", [
    ("150000", "450000", "80", True),
    ("0", "450000", "80", False),        # limit 0 olamaz
    ("150000", "-1", "80", False),       # negatif limit
    ("150000", "450000", "0", False),    # yüzde 0 olamaz
    ("150000", "450000", "120", False),  # yüzde 100'den büyük olamaz
    ("abc", "450000", "80", False),      # sayı değil
    ("450000", "150000", "80", True),    # tek>toplam teknik olarak geçerli (engellenmez)
])
def test_kriter_dogrula(et, eto, y, ok):
    assert exay.kriter_dogrula(et, eto, y)[0] is ok


# ══════════════════════════════════════════════════════════════════════════
#  Doğruluk kontrolleri (uyarı üreticiler)
# ══════════════════════════════════════════════════════════════════════════
def test_kdv_tutarlilik_kotu_esleme_uyarir():
    # KDV = matrah (oran %100) → şüpheli, uyarı beklenir
    df = pd.DataFrame({
        "Alış Faturasının KDV Hariç Tutarı": [1000, 2000],
        "KDV si": [1000, 2000],
    })
    assert exay.kdv_tutarlilik_kontrol(df)          # boş değil (uyarı var)


def test_kdv_tutarlilik_iyi_esleme_uyarmaz():
    df = pd.DataFrame({
        "Alış Faturasının KDV Hariç Tutarı": [1000, 2000],
        "KDV si": [180, 360],                        # %18
    })
    assert exay.kdv_tutarlilik_kontrol(df) == []


def test_mukerrer_fatura():
    df = pd.DataFrame({
        "Satıcının Vergi Kimlik Numarası": ["1000000001", "1000000001", "1000000002"],
        "Alış Faturasının Sıra No'su": ["F1", "F1", "F2"],
    })
    m = exay.mukerrer_fatura_bul(df)
    assert m == [("1000000001", "F1", 2)]


def test_ay_yil_bicimleri():
    from datetime import datetime as _dt
    assert exay._ay_yil("2026-04-15") == "04.2026"    # ISO
    assert exay._ay_yil("15.04.2026") == "04.2026"    # gün.ay.yıl
    assert exay._ay_yil(_dt(2026, 4, 15)) == "04.2026"
    assert exay._ay_yil("") is None


def test_donem_disi_tarih():
    df = pd.DataFrame({
        "Alış Faturasının Tarihi": ["2026-04-01", "2026-04-02", "2026-05-15"],
    })
    toplam, disi = exay.donem_disi_tarih_kontrol(df, "04.2026")
    assert toplam == 3 and disi == 1                  # yalnızca Mayıs dönem dışı


# ══════════════════════════════════════════════════════════════════════════
#  Sürükle-bırak yolu ayıklama
# ══════════════════════════════════════════════════════════════════════════
def test_dnd_ayikla():
    ayik = exay.KDVBolmeApp._dnd_ayikla
    assert ayik("/tmp/a.xlsx") == ["/tmp/a.xlsx"]
    assert ayik("{/tmp/bir iki.xlsx} /tmp/c.csv") == ["/tmp/bir iki.xlsx", "/tmp/c.csv"]


# ══════════════════════════════════════════════════════════════════════════
#  PDF çıktı (reportlab kuruluysa)
# ══════════════════════════════════════════════════════════════════════════
@pytest.mark.skipif(not exay.pdf_destekli(), reason="reportlab kurulu değil")
def test_pdf_uretimi(tmp_path):
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)
    df = exay.ana_listeyi_oku(str(yol))
    secilen, _ = exay.firmalari_filtrele(df, 150000, 450000, 80, _sessiz)
    grp = secilen["1000000001"][0]
    pdf = tmp_path / "firma.pdf"
    exay.firma_pdf_olustur(grp, str(pdf), list(df.columns),
                           vkn="1000000001", unvan="ÇĞİÖŞÜ FİRMA", donem="04.2026")
    assert pdf.exists() and pdf.read_bytes()[:5] == b"%PDF-"


def test_dosyalari_isle_pdf_uret(tmp_path):
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        pdf_uret=True)
    cikis = tmp_path / "Hazır Tutanaklar"
    pdfler = list(cikis.glob("*.pdf"))
    if exay.pdf_destekli():
        assert len(pdfler) == 2                       # A ve B için PDF
    else:
        assert len(pdfler) == 0                       # reportlab yoksa sessizce atlanır


# ══════════════════════════════════════════════════════════════════════════
#  Çıktı klasörü yönlendirme (cikis_kok)
# ══════════════════════════════════════════════════════════════════════════
def test_cikis_kok_yonlendirme(tmp_path):
    kaynak = tmp_path / "kaynak"
    kaynak.mkdir()
    yol = kaynak / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)
    hedef = tmp_path / "baska_yer"
    hedef.mkdir()
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        cikis_kok=str(hedef))
    # Çıktı kaynağın yanına DEĞİL, seçilen hedefe yazılmalı
    assert (hedef / "Hazır Tutanaklar").is_dir()
    assert not (kaynak / "Hazır Tutanaklar").exists()


# ══════════════════════════════════════════════════════════════════════════
#  WORD ŞABLON EŞLEŞTİRME (VKN ile) — saf ayrıştırıcılar
# ══════════════════════════════════════════════════════════════════════════
def _tutanak_metni(unvan, vd_hucre):
    """Gerçek tutanak düzenini (sekmeli hücreler) taklit eden sentetik metin."""
    return (
        "KATMA DEĞER VERGİSİ İADESİ KARŞIT İNCELEME TUTANAĞI\n"
        "YEMİNLİ MALİ MÜŞAVİRİN\t\tAdı Soyadı\tSABRİ HAMAMCI\t\t"
        "Vergi Dairesi ve Sicil No\tGAZİKENT V.D. / 464 100 8244\t\t"
        "İADE TALEBİNDE BULUNAN FİRMANIN\t\tÜnvanı\tİNALOĞLU İNŞAAT\t\t"
        "Vergi Dairesi/Nosu\tŞAHİNBEY / 475 056 9431\t\t"
        "NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN\t\t"
        f"Ünvanı\t{unvan}\t\tVergi Dairesi/Nosu\t{vd_hucre}\t\t"
        "Adresi\tBİR ADRES\t\tİNCELEME DAYANAĞI\t31.01.2026 Tarih ve 09 Sayılı\n"
        "Karşıt İncelemeye Konu Fatura ve Benzeri Belgeye ilişkin bilgiler:\n"
    )


def test_vkn_metinden_ayikla():
    assert exay._vkn_metinden_ayikla("493 061 9102") == "4930619102"   # boşluklu
    assert exay._vkn_metinden_ayikla("ASIM GÜNDÜZ V.D. – 30490690382") == "30490690382"
    assert exay._vkn_metinden_ayikla("71419747") == "0071419747"       # 8 hane → zfill
    assert exay._vkn_metinden_ayikla("0000000000") is None             # yer tutucu
    assert exay._vkn_metinden_ayikla("yok") is None


def test_sablon_vkn_metinden():
    # 10 haneli VKN (boşluklu) doğru bloktan alınmalı, ünvan da
    vkn, unvan = exay.sablon_vkn_metinden(
        _tutanak_metni("MESAKO MADEN VE ENERJİ TİC. LTD. ŞTİ.", "ŞAHİNBEY V.D. 6190914983"))
    assert vkn == "6190914983"
    assert "MESAKO" in unvan
    # 11 haneli TCKN + tire ile
    vkn2, unvan2 = exay.sablon_vkn_metinden(
        _tutanak_metni("ONUR FURKAN KARTA", "ASIM GÜNDÜZ V.D. – 30490690382"))
    assert vkn2 == "30490690382"
    assert unvan2 == "ONUR FURKAN KARTA"
    # Blok yoksa (ör. üst yazı) → (None, None)
    assert exay.sablon_vkn_metinden("Sayı: YMM 27103572 ... GAZİANTEP") == (None, None)


def test_sablonlari_indeksle(tmp_path, monkeypatch):
    metinler = {}
    for ad, unvan, vkn in [("a.doc", "FIRMA A", "1234567890"),
                           ("b.doc", "FIRMA B", "9876543210"),
                           ("ustyazi.doc", None, None)]:
        p = tmp_path / ad
        p.write_bytes(b"stub")     # gerçek .doc gerekmez; okuma monkeypatch'li
        if unvan:
            metinler[str(p)] = _tutanak_metni(unvan, f"V.D. {vkn}")
        else:
            metinler[str(p)] = "Sayı: YMM 123 üst yazı"
    monkeypatch.setattr(exay, "_doc_metni_oku", lambda p: metinler[str(p)])
    idx = exay.sablonlari_indeksle(str(tmp_path))
    assert idx["1234567890"][0].endswith("a.doc")
    assert idx["9876543210"][0].endswith("b.doc")
    assert len(idx) == 2               # üst yazı eşleşmez


def test_fatura_tablosu_mu():
    assert exay._fatura_tablosu_mu(["FATURANIN", "MALIN", "Tarihi", "Numarası",
                                    "Cinsi", "Miktarı", "Tutarı", "KDV Tutarı"])
    assert not exay._fatura_tablosu_mu(["Defterin Nevi", "Tasdik Makamı"])


def test_word_fatura_satiri_ve_miktar():
    df = pd.DataFrame({
        "Alış Faturasının Tarihi": ["2026-04-01"],
        "Alış Faturasının Sıra No'su": ["BBK1"],
        "Alınan Mal ve/veya Hizmetin Cinsi": ["YEMEK BEDELİ"],
        "Miktarı": ["4.880 ADET"],
        "Alış Faturasının KDV Hariç Tutarı": [683200],
        "KDV si": [68320],
    })
    hucre = exay._word_fatura_satiri(df.iloc[0], list(df.columns))
    # [Tarih, No, Cinsi, Miktar, Tutar, KDV, DefterKayıt(boş)]
    assert hucre[0] == "01.04.2026"
    assert hucre[1] == "BBK1"
    assert hucre[2] == "YEMEK BEDELİ"
    assert hucre[3] == "4.880 ADET"          # miktar sütunu kullanıldı
    assert hucre[4] == "683.200,00"          # TR biçim
    assert hucre[5] == "68.320,00"
    assert hucre[6] == ""                     # defter kayıt boş


def test_tr_para_str():
    assert exay._tr_para_str(683200) == "683.200,00"
    assert exay._tr_para_str("1234567,89") == "1.234.567,89"
    assert exay._tr_para_str("") == ""


def test_word_com_yoksa_hata():
    # Bu ortamda Word/pywin32 yok → firma_word_olustur RuntimeError vermeli
    if exay.word_destekli():
        pytest.skip("Word otomasyonu mevcut; negatif yol test edilemez")
    df = pd.DataFrame({"Alış Faturasının Tarihi": ["2026-04-01"]})
    with pytest.raises(RuntimeError):
        exay.firma_word_olustur("yok.doc", df, "cikti.doc", list(df.columns))


def test_dosyalari_isle_sablon_eslesme_raporu(tmp_path, monkeypatch):
    """Şablon klasörü verildiğinde: eşleşme yapılır ve WORD_ESLESME raporu üretilir
    (Word olmadan da). Eşleşme VKN ile olmalı."""
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)     # FIRMA A=1000000001 (seçilir), FIRMA B=1000000002 (seçilir)
    sablon_kl = tmp_path / "sablonlar"
    sablon_kl.mkdir()
    pa = sablon_kl / "firmaA.doc"; pa.write_bytes(b"stub")
    metin = _tutanak_metni("FIRMA A", "V.D. 1000000001")   # yalnızca A'nın şablonu
    monkeypatch.setattr(exay, "_doc_metni_oku",
                        lambda p: metin if str(p) == str(pa) else "üst yazı")
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sablon_kl))
    rapor = list((tmp_path / "Hazır Tutanaklar").glob("WORD_ESLESME_*.xlsx"))
    assert rapor, "Şablon eşleşme raporu üretilmedi"
    wb = openpyxl.load_workbook(rapor[0]); ws = wb.active
    satirlar = {ws.cell(r, 1).value: ws.cell(r, 3).value for r in range(2, ws.max_row + 1)}
    assert "1000000001" in satirlar             # A eşleşti
    assert "1000000002" in satirlar             # B şablonsuz
    assert satirlar["1000000002"] == "Şablon yok"


def test_gercek_gib_basliklari_word_eslemesi():
    """Gerçek 'Yeni GİB' başlık düzeni (uzun sütun adları, kesme işaretli KDV'si,
    birleşik VKN/TC sütunu, 'Alınan Mal ve/veya Hizmetin Miktarı') doğru eşlenmeli
    ve Word satırı [Tarih,No,Cinsi,Miktar,Tutar,KDV,''] üretmeli."""
    kols = [
        "Alış Faturasının Tarihi", "Alış Faturasının Serisi",
        "Alış Faturasının Sıra No'su", "Satıcının Adı-Soyadı / Ünvanı",
        "Satıcının Vergi Kimlik Numarası / TC Kimlik Numarası",
        "Alınan Mal ve/veya Hizmetin Cinsi", "Alınan Mal ve/veya Hizmetin Miktarı",
        "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı", "KDV'si",
        "Toplam İndirilecek KDV Tutarı",
    ]
    df = pd.DataFrame([[
        "2026-07-25", None, "0012026165876119", "TURKCELL ILETİSİM HİZMETLERİ A.S.",
        "8770013456", "İLETİŞİM HİZMET BEDELİ", "5 ADET", 1208.34, 241.67, 241.67,
    ]], columns=kols)
    b = exay.bulunan_sutunlar(df)
    assert b["Matrah"] == "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı"
    assert b["KDV"] == "KDV'si"                          # kesme işaretli, toplam KDV ile karışmaz
    assert exay.sutun_bul(kols, ["miktar"]) == "Alınan Mal ve/veya Hizmetin Miktarı"
    satir = exay._word_fatura_satiri(df.iloc[0], kols)
    assert satir == ["25.07.2026", "0012026165876119", "İLETİŞİM HİZMET BEDELİ",
                     "5 ADET", "1.208,34", "241,67", ""]


# ══════════════════════════════════════════════════════════════════════════
#  GUI kurulum dumanı — _ui() sahte pencereyle çalışır; eksik metot/callback
#  bağlamaları (ör. bind command'ları) burada yakalanır (regresyon güvencesi).
# ══════════════════════════════════════════════════════════════════════════
def test_gui_kurulur_ve_callbackler_tanimli():
    from unittest.mock import MagicMock
    root = MagicMock(name="root")
    app = exay.KDVBolmeApp(root)     # __init__ → _ui() çalışır; eksik metot patlar
    # _ui / _surukle_birak içinde referans verilen tüm callback'ler tanımlı olmalı
    for m in ["_tiklayarak_sec", "_isle", "_isle_coklu", "_batch_worker",
              "_cikis_klasoru_sec", "_sablon_klasoru_sec", "_cikis_ozet",
              "_sablon_ozet", "_kriter_al", "_ilerleme", "_tamam", "_log",
              "_ayar_kaydet", "_birak_guncelle", "_dnd_ayikla"]:
        assert callable(getattr(app, m)), f"Eksik/çağrılamaz metot: {m}"


# ══════════════════════════════════════════════════════════════════════════
#  .docx şablon yolu — okuma (VKN) + yazma (fatura tablosu), Word GEREKTİRMEZ
# ══════════════════════════════════════════════════════════════════════════
def _docx_sablon_yaz(yol, unvan, vd_hucre):
    """Gerçek tutanak yapısını taklit eden sentetik .docx şablon üretir:
    2 sütunlu bilgi tablosu (NEZDİNDE bloğu) + 7 sütunlu fatura tablosu."""
    import docx
    d = docx.Document()
    t0 = d.add_table(rows=0, cols=2)
    for etiket, deger in [
        ("YEMİNLİ MALİ MÜŞAVİRİN", "SABRİ HAMAMCI"),
        ("İADE TALEBİNDE BULUNAN FİRMANIN", "İADE TALEBİNDE BULUNAN FİRMANIN"),
        ("Ünvanı", "İNALOĞLU İNŞAAT"),
        ("Vergi Dairesi/Nosu", "ŞAHİNBEY / 475 056 9431"),
        ("NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN", "NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN"),
        ("Ünvanı", unvan),
        ("Vergi Dairesi/Nosu", vd_hucre),
        ("İNCELEME DAYANAĞI", "31.01.2026 Tarih ve 09 Sayılı"),
    ]:
        r = t0.add_row().cells
        r[0].text = etiket; r[1].text = deger
    d.add_paragraph("Karşıt İncelemeye Konu Fatura ve Benzeri Belgeye ilişkin bilgiler:")
    t2 = d.add_table(rows=3, cols=7)
    b0 = ["FATURANIN", "FATURANIN", "MALIN", "MALIN", "MALIN", "MALIN", "Defter Kayıt"]
    b1 = ["Tarihi", "Numarası", "Cinsi", "Miktarı", "Tutarı", "KDV Tutarı", "Tarihi/Nosu"]
    for c, v in enumerate(b0): t2.rows[0].cells[c].text = v
    for c, v in enumerate(b1): t2.rows[1].cells[c].text = v
    d.save(yol)


def test_docx_sablon_vkn_oku(tmp_path):
    yol = tmp_path / "ispa.docx"
    _docx_sablon_yaz(yol, "İSPA İNŞ. SAN. PAZ. A.Ş.", "ÜSKÜDAR V.D. – 481 001 7371")
    vkn, unvan = exay.sablon_vkn_oku(str(yol))
    assert vkn == "4810017371"                       # boşluklu VKN düzeltilir
    assert "İSPA" in unvan


def test_sablonlari_indeksle_docx_dahil(tmp_path):
    _docx_sablon_yaz(tmp_path / "a.docx", "FIRMA A", "V.D. 1234567890")
    idx = exay.sablonlari_indeksle(str(tmp_path))
    assert idx["1234567890"][0].endswith("a.docx")   # .docx da indekslenir


def test_firma_docx_olustur_tabloyu_doldurur(tmp_path):
    import docx
    sablon = tmp_path / "sablon.docx"
    _docx_sablon_yaz(sablon, "İSPA İNŞ. SAN. PAZ. A.Ş.", "V.D. 4810017371")
    # iki faturalı bir firma
    kols = ["Alış Faturasının Tarihi", "Alış Faturasının Sıra No'su",
            "Alınan Mal ve/veya Hizmetin Cinsi", "Alınan Mal ve/veya Hizmetin Miktarı",
            "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı", "KDV'si"]
    firma = pd.DataFrame([
        ["2026-07-22", "S0120260058", "İNŞAAT MALZ.", "6 ADET", 2595795.15, 516799.39],
        ["2026-07-25", "S0120260059", "DEMİR", "10 TON", 100000.00, 20000.00],
    ], columns=kols)
    cikti = tmp_path / "cikti.docx"
    exay.firma_docx_olustur(str(sablon), firma, str(cikti), kols)
    d = docx.Document(str(cikti))
    fatura = next(t for t in d.tables if len(t.columns) == 7)
    # 2 başlık satırı + 2 veri satırı = 4
    assert len(fatura.rows) == 4
    veri = [[c.text.strip() for c in r.cells] for r in fatura.rows[2:]]
    assert veri[0][0] == "22.07.2026"
    assert veri[0][1] == "S0120260058"
    assert veri[0][3] == "6 ADET"                    # miktar
    assert veri[0][4] == "2.595.795,15"              # TR tutar
    assert veri[0][5] == "516.799,39"
    assert veri[1][2] == "DEMİR"
    assert veri[1][3] == "10 TON"


def test_dosyalari_isle_docx_word_uretir(tmp_path):
    """Uçtan uca: .docx şablon eşleşince gerçek Word tutanağı üretilmeli (Word yok)."""
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)     # FIRMA A = 1000000001 (seçilir)
    sk = tmp_path / "sablonlar"; sk.mkdir()
    _docx_sablon_yaz(sk / "firmaA.docx", "FIRMA A", "V.D. 1000000001")
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk))
    uretilen = list((tmp_path / "Hazır Tutanaklar").glob("*.docx"))
    assert any("1000000001" in p.name for p in uretilen), "İSPA benzeri Word tutanağı üretilmedi"


# ══════════════════════════════════════════════════════════════════════════
#  Çıktı türü seçimi: yalnız Excel / yalnız Word / ikisi
# ══════════════════════════════════════════════════════════════════════════
def _liste_ve_sablon(tmp_path):
    """FIRMA A (1000000001) seçilir; A için bir .docx şablon hazırlanır."""
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)
    sk = tmp_path / "sablonlar"; sk.mkdir()
    _docx_sablon_yaz(sk / "firmaA.docx", "FIRMA A", "V.D. 1000000001")
    return yol, sk


def test_cikti_yalniz_excel(tmp_path):
    yol, sk = _liste_ve_sablon(tmp_path)
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk), cikti_turu='excel')
    kl = tmp_path / "Hazır Tutanaklar"
    assert list(kl.glob("*.xlsx"))                       # Excel var
    assert not list(kl.glob("*.docx"))                    # Word YOK
    assert not list(kl.glob("WORD_ESLESME_*.xlsx"))       # eşleşme raporu bile yok


def test_cikti_yalniz_word(tmp_path):
    yol, sk = _liste_ve_sablon(tmp_path)
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk), cikti_turu='word')
    kl = tmp_path / "Hazır Tutanaklar"
    docx_ler = list(kl.glob("*.docx"))
    # Firma tutanağı .docx üretilir; VKN listesi/özet .xlsx yan dosyaları hariç
    # FİRMA tutanağı .xlsx OLMAMALI (numaralı 'N) …_.xlsx')
    firma_xlsx = [p for p in kl.glob("*.xlsx")
                  if p.name[0].isdigit() and ")" in p.name[:4]]
    assert any("1000000001" in p.name for p in docx_ler)  # A için Word üretildi
    assert not firma_xlsx                                  # firma Excel tutanağı YOK


def test_cikti_ikisi(tmp_path):
    yol, sk = _liste_ve_sablon(tmp_path)
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk), cikti_turu='ikisi')
    kl = tmp_path / "Hazır Tutanaklar"
    firma_xlsx = [p for p in kl.glob("*.xlsx")
                  if p.name[0].isdigit() and ")" in p.name[:4]]
    assert firma_xlsx                                      # Excel tutanakları var (A ve B)
    assert any("1000000001" in p.name for p in kl.glob("*.docx"))  # A için Word da var


# ══════════════════════════════════════════════════════════════════════════
#  İnceleme Dayanağı (sözleşme) otomatik güncelleme + Türkçe casing güvenliği
# ══════════════════════════════════════════════════════════════════════════
def test_ascii_kucuk_turkce():
    assert exay._ascii_kucuk("İNCELEME DAYANAĞI") == "inceleme dayanagi"
    assert "inceleme dayana" in exay._ascii_kucuk("İNCELEME DAYANAĞI")
    assert exay._ascii_kucuk("ŞUBAT") == "subat"


def test_inceleme_dayanagi_docx_gunceller(tmp_path):
    import docx
    sablon = tmp_path / "s.docx"
    _docx_sablon_yaz(sablon, "İSPA A.Ş.", "V.D. 4810017371")
    kols = ["Alış Faturasının Tarihi", "Alış Faturasının Sıra No'su",
            "Alınan Mal ve/veya Hizmetin Cinsi", "Alınan Mal ve/veya Hizmetin Miktarı",
            "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı", "KDV'si"]
    firma = pd.DataFrame([["2026-07-22", "F1", "MAL", "6 ADET", 1000.0, 180.0]], columns=kols)
    cikti = tmp_path / "c.docx"
    YENI = "30.06.2026 Tarih ve 27 Sayılı Tam Tasdik Sözleşmesi"
    exay.firma_docx_olustur(str(sablon), firma, str(cikti), kols, inceleme_dayanagi=YENI)
    d = docx.Document(str(cikti))
    bulundu = False
    for tbl in d.tables:
        for row in tbl.rows:
            cs = row.cells
            if cs and 'inceleme dayana' in exay._ascii_kucuk(cs[0].text):
                assert cs[1].text.strip() == YENI          # değer güncellendi
                assert 'inceleme dayana' in exay._ascii_kucuk(cs[0].text)  # etiket sabit
                bulundu = True
    assert bulundu


# ══════════════════════════════════════════════════════════════════════════
#  Çok-firmalı tek .docx: içinden ilgili firmayı VKN ile bulup izole etme
# ══════════════════════════════════════════════════════════════════════════
def _docx_coklu_yaz(yol, firmalar):
    """Tek dosyada birden çok firma tutanağı üretir; her blok
    'KATMA DEĞER ... KARŞIT İNCELEME TUTANAĞI' başlığıyla başlar (gerçek yapı)."""
    import docx
    d = docx.Document()
    for k, (unvan, vd) in enumerate(firmalar):
        if k > 0:
            d.add_page_break()
        d.add_paragraph("KATMA DEĞER VERGİSİ İADESİ KARŞIT İNCELEME TUTANAĞI")
        t0 = d.add_table(rows=0, cols=2)
        for e, v in [("NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN", e := "NEZDİNDE KARŞIT İNCELEME YAPILAN FİRMANIN"),
                     ("Ünvanı", unvan), ("Vergi Dairesi/Nosu", vd),
                     ("İNCELEME DAYANAĞI", "31.01.2026 Tarih ve 09 Sayılı")]:
            r = t0.add_row().cells; r[0].text = e if isinstance(e, str) else e; r[1].text = v
        d.add_paragraph("Karşıt İncelemeye Konu Fatura ve Benzeri Belgeye ilişkin bilgiler:")
        t2 = d.add_table(rows=3, cols=7)
        for c, v in enumerate(["FATURANIN", "FATURANIN", "MALIN", "MALIN", "MALIN", "MALIN", "Defter Kayıt"]):
            t2.rows[0].cells[c].text = v
        for c, v in enumerate(["Tarihi", "Numarası", "Cinsi", "Miktarı", "Tutarı", "KDV Tutarı", "Tarihi/Nosu"]):
            t2.rows[1].cells[c].text = v
    d.save(yol)


def test_docx_coklu_firma_kayitlari(tmp_path):
    yol = tmp_path / "hepsi.docx"
    _docx_coklu_yaz(yol, [("FIRMA A", "V.D. 1234567890"),
                          ("FIRMA B", "V.D. 9876543210")])
    kayit = exay._sablon_kayitlari(str(yol))
    vknler = {v: b for v, u, b, y in kayit}
    assert vknler == {"1234567890": 0, "9876543210": 1}   # iki firma, blok 0/1
    idx = exay.sablonlari_indeksle(str(tmp_path))
    assert idx["1234567890"] == (str(yol), 0)
    assert idx["9876543210"] == (str(yol), 1)


def test_docx_blok_izole_ve_doldur(tmp_path):
    import docx
    yol = tmp_path / "hepsi.docx"
    _docx_coklu_yaz(yol, [("FIRMA A", "V.D. 1234567890"),
                          ("FIRMA B", "V.D. 9876543210")])
    kols = ["Alış Faturasının Tarihi", "Alış Faturasının Sıra No'su",
            "Alınan Mal ve/veya Hizmetin Cinsi", "Alınan Mal ve/veya Hizmetin Miktarı",
            "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı", "KDV'si"]
    firma = pd.DataFrame([["2026-07-02", "B1", "MAL", "5 ADET", 2000.0, 360.0]], columns=kols)
    cikti = tmp_path / "b.docx"
    # blok 1 = FIRMA B izole edilmeli
    exay.firma_docx_olustur(str(yol), firma, str(cikti), kols, blok=1)
    d = docx.Document(str(cikti))
    # Tek firma → sadece B'nin bloğu (1 fatura tablosu) kalmalı
    fat = [t for t in d.tables if len(t.columns) == 7]
    assert len(fat) == 1
    # içindeki firma B olmalı (A değil)
    metin = "\n".join("\t".join(c.text.strip() for c in r.cells)
                      for t in d.tables for r in t.rows)
    vkn, _ = exay.sablon_vkn_metinden(metin)
    assert vkn == "9876543210"
    # fatura satırı dolduruldu
    veri = [c.text.strip() for c in fat[0].rows[2].cells]
    assert veri[1] == "B1" and veri[3] == "5 ADET"


def test_docx_coklu_uctan_uca(tmp_path):
    """Liste + çok-firmalı tek şablon dosyası → her firma ayrı tutanak."""
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)     # FIRMA A=1000000001, FIRMA B=1000000002 seçilir
    sk = tmp_path / "sablonlar"; sk.mkdir()
    _docx_coklu_yaz(sk / "hepsi.docx", [("FIRMA A", "V.D. 1000000001"),
                                        ("FIRMA B", "V.D. 1000000002")])
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk), cikti_turu='word')
    uretilen = sorted(p.name for p in (tmp_path / "Hazır Tutanaklar").glob("*.docx"))
    assert any("1000000001" in n for n in uretilen)   # A tek dosyadan bulundu
    assert any("1000000002" in n for n in uretilen)   # B tek dosyadan bulundu


# ══════════════════════════════════════════════════════════════════════════
#  Tek dosyada birleştirme (word_tek_dosya) + boş/yedek şablon (bos_sablon)
# ══════════════════════════════════════════════════════════════════════════
def test_word_tek_dosya_birlestirme(tmp_path):
    import docx
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)   # A=1000000001, B=1000000002 seçilir
    sk = tmp_path / "sablonlar"; sk.mkdir()
    _docx_sablon_yaz(sk / "a.docx", "FIRMA A", "V.D. 1000000001")
    _docx_sablon_yaz(sk / "b.docx", "FIRMA B", "V.D. 1000000002")
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk), cikti_turu='word', word_tek_dosya=True)
    kl = tmp_path / "Hazır Tutanaklar"
    birlesik = list(kl.glob("KARSIT_INCELEME_TUTANAKLAR_*.docx"))
    assert len(birlesik) == 1                              # tek birleşik dosya
    # ayrı firma .docx'i OLMAMALI
    assert not [p for p in kl.glob("*.docx") if p.name[0].isdigit()]
    d = docx.Document(str(birlesik[0]))
    fat = [t for t in d.tables if len(t.columns) == 7]
    assert len(fat) == 2                                   # A ve B blokları tek dosyada


def test_bos_sablon_eslesmeyen_firma(tmp_path):
    import docx
    yol = tmp_path / "NISAN_2026.xlsx"
    _yeni_gib_yaz(yol)   # A=1000000001, B=1000000002
    sk = tmp_path / "sablonlar"; sk.mkdir()
    _docx_sablon_yaz(sk / "a.docx", "FIRMA A", "V.D. 1000000001")   # yalnızca A'nın şablonu
    bos = tmp_path / "bos.docx"
    _docx_sablon_yaz(bos, "", "")                          # boş NEZDİNDE (ünvan/vd yok)
    exay.dosyalari_isle(str(yol), 150000, 450000, 80, _sessiz, lambda *a: None,
                        sablon_klasor=str(sk), cikti_turu='word', bos_sablon=str(bos))
    kl = tmp_path / "Hazır Tutanaklar"
    # B için boş şablondan üretilmiş bir dosya olmalı
    b_dosya = [p for p in kl.glob("*.docx") if "1000000002" in p.name]
    assert b_dosya, "Boş şablondan B tutanağı üretilmedi"
    d = docx.Document(str(b_dosya[0]))
    metin = "\n".join("\t".join(c.text.strip() for c in r.cells)
                      for t in d.tables for r in t.rows)
    vkn, unvan = exay.sablon_vkn_metinden(metin)
    assert vkn == "1000000002"                            # bilinen VKN yazıldı
    assert unvan == "FIRMA B"                             # bilinen ünvan yazıldı
    # WORD_ESLESME raporunda "Boş şablon" durumu geçmeli
    rap = list(kl.glob("WORD_ESLESME_*.xlsx"))[0]
    wb = openpyxl.load_workbook(rap); ws = wb.active
    durumlar = {ws.cell(r, 1).value: ws.cell(r, 3).value for r in range(2, ws.max_row + 1)}
    assert "Boş şablon" in str(durumlar.get("1000000002", ""))


# ══════════════════════════════════════════════════════════════════════════
#  Birleşik .doc tanıma (metin bölme) — üretim COM gerektirir, burada tanıma test
# ══════════════════════════════════════════════════════════════════════════
def test_metni_bloklara_ayir():
    tek = _tutanak_metni("FIRMA A", "V.D. 1234567890")
    assert len(exay._metni_bloklara_ayir(tek)) == 1
    ikili = _tutanak_metni("FIRMA A", "V.D. 1234567890") + \
            _tutanak_metni("FIRMA B", "V.D. 9876543210")
    bloklar = exay._metni_bloklara_ayir(ikili)
    assert len(bloklar) == 2
    assert exay.sablon_vkn_metinden(bloklar[0])[0] == "1234567890"
    assert exay.sablon_vkn_metinden(bloklar[1])[0] == "9876543210"


def test_sablon_kayitlari_doc_tekli(monkeypatch, tmp_path):
    p = tmp_path / "tek.doc"; p.write_bytes(b"stub")
    monkeypatch.setattr(exay, "_doc_metni_oku",
                        lambda x: _tutanak_metni("FIRMA A", "V.D. 1234567890"))
    kayit = exay._sablon_kayitlari(str(p))
    assert kayit == [("1234567890", kayit[0][1], None, str(p))]   # tek firma, blok None


# ══════════════════════════════════════════════════════════════════════════
#  İkinci belge tipi: YMM 'Bilgi İsteme' yazısı (Hakkında Bilgi İstenilen Mükellef)
# ══════════════════════════════════════════════════════════════════════════
def test_ymm_yazisi_vkn_eslesme():
    # Etiketler karışık olabilir (Adresi hücresinde V.D./VKN) ve telefon var
    metin = (
        "Sayı : YMM 27103572/2026-363\tGAZİANTEP\n"
        "İade Talebinde Bulunan Firma\t\tUnvanı\tİNALOĞLU İNŞAAT\t\t"
        "Adresi\tŞAHİNBEY / 475 056 9431\t\tTelefon/Fax\t0 342 502 03 15\n"
        "Hakkında Bilgi İstenilen Mükellefin Altı\t\t"
        "Ünvanı\tOYAK ÇİMENTO FABRİKALARI ANONİM ŞİRKETİ\t\t"
        "Adresi\tANKARA KURUMLAR V.D. – 6120050961\t\t"
        "Vergi Dairesi/Hesap Nosu\tÇUKURAMBAR MAH. 1480 SK.\t\t"
        "Telefon/Fax\t0 312 220 0290\n"
        "İNCELEME DAYANAĞI\t31.01.2026 Tarih ve 09 Sayılı\n"
    )
    vkn, unvan = exay.sablon_vkn_metinden(metin)
    assert vkn == "6120050961"          # karşı firma (telefon 03122200290 DEĞİL)
    assert "OYAK" in unvan


def test_blok_vkn_telefon_karistirmaz():
    blok = "Ünvanı\tX A.Ş.\tTelefon/Fax\t0 342 215 10 70\tVergi Dairesi\tŞAHİNBEY V.D. 6190914983"
    assert exay._blok_vkn(blok) == "6190914983"


# ── YMM yazısı: fatura tablosunun son sütunu 'KDV dahil toplam' (tutanakta boş) ──
def _ymm_yazi_docx_yaz(yol, unvan, vd_hucre):
    """Gerçek YMM 'Bilgi İsteme' yazısını taklit eden sentetik .docx: 'Hakkında
    Bilgi İstenilen Mükellef' bloğu + 7 sütunlu, son sütunu 'KDV dahil toplam'
    olan (tek başlık satırlı) fatura tablosu."""
    import docx
    d = docx.Document()
    d.add_paragraph("Sayı : YMM 27103572/2026-363")
    d.add_paragraph("Konu : Bilgi İsteme")
    t0 = d.add_table(rows=0, cols=2)
    for etiket, deger in [
        ("Hakkında Bilgi İstenilen Mükellefin", "Hakkında Bilgi İstenilen Mükellefin"),
        ("Ünvanı", unvan),
        ("Vergi Dairesi/Hesap Nosu", vd_hucre),
        ("Telefon/Fax", "0 312 220 0290"),
        ("İNCELEME DAYANAĞI", "31.01.2026 Tarih ve 09 Sayılı"),
    ]:
        r = t0.add_row().cells
        r[0].text = etiket; r[1].text = deger
    t2 = d.add_table(rows=1, cols=7)
    bas = ["FAT.TARİHİ", "FAT. NOSU", "MALIN CİNSİ", "MALIN MİKTARI",
           "MATRAH", "KDV", "kdv dahİl toplam"]
    for c, v in enumerate(bas):
        t2.rows[0].cells[c].text = v
    d.save(yol)


def test_fatura_kolon_rolu_iki_tip():
    # Tutanak: son sütun 'Defter Kayıt' → boş; YMM: 'KDV dahil toplam' → dahil
    assert exay._fatura_kolon_rolu("Defter Kayıt Tarihi/Nosu") == "bos"
    assert exay._fatura_kolon_rolu("kdv dahİl toplam") == "dahil"
    assert exay._fatura_kolon_rolu("FATURANIN Tarihi") == "tarih"
    assert exay._fatura_kolon_rolu("FAT. NOSU") == "no"
    assert exay._fatura_kolon_rolu("MALIN KDV Tutarı") == "kdv"
    assert exay._fatura_kolon_rolu("MATRAH") == "matrah"
    assert exay._fatura_kolon_rolu("MALIN Miktarı") == "miktar"


def _ornek_firma_df():
    kols = ["Alış Faturasının Tarihi", "Alış Faturasının Sıra No'su",
            "Alınan Mal ve/veya Hizmetin Cinsi", "Alınan Mal ve/veya Hizmetin Miktarı",
            "Alınan Mal ve/veya Hizmetin KDV Hariç Tutarı", "KDV'si"]
    firma = pd.DataFrame([
        ["2026-01-08", "TC42026000001608", "HAZIR BETON", "302 M3", 446641.18, 89328.24],
    ], columns=kols)
    return firma, kols


def test_ymm_yazi_fatura_kdv_dahil_toplam(tmp_path):
    import docx
    sablon = tmp_path / "ymm.docx"
    _ymm_yazi_docx_yaz(sablon, "OYAK ÇİMENTO FABRİKALARI A.Ş.", "ANKARA KURUMLAR V.D. – 6120050961")
    firma, kols = _ornek_firma_df()
    cikti = tmp_path / "ymm_cikti.docx"
    exay.firma_docx_olustur(str(sablon), firma, str(cikti), kols)
    d = docx.Document(str(cikti))
    fatura = next(t for t in d.tables if len(t.columns) == 7)
    veri = [c.text.strip() for c in fatura.rows[1].cells]      # tek başlık satırı
    assert veri[0] == "08.01.2026"
    assert veri[4] == "446.641,18"                              # matrah
    assert veri[5] == "89.328,24"                               # kdv
    assert veri[6] == "535.969,42"                              # KDV dahil toplam = matrah+kdv


def test_tutanak_defter_kayit_bos(tmp_path):
    import docx
    sablon = tmp_path / "tut.docx"
    _docx_sablon_yaz(sablon, "İSPA A.Ş.", "V.D. 4810017371")
    firma, kols = _ornek_firma_df()
    cikti = tmp_path / "tut_cikti.docx"
    exay.firma_docx_olustur(str(sablon), firma, str(cikti), kols)
    d = docx.Document(str(cikti))
    fatura = next(t for t in d.tables if len(t.columns) == 7)
    veri = [c.text.strip() for c in fatura.rows[2].cells]      # 2 başlık satırı
    assert veri[4] == "446.641,18" and veri[5] == "89.328,24"
    assert veri[6] == ""                                        # Defter Kayıt boş kalır


def test_ymm_yazi_vkn_indekslenir(tmp_path):
    _ymm_yazi_docx_yaz(tmp_path / "oyak.docx", "OYAK ÇİMENTO A.Ş.",
                       "ANKARA KURUMLAR V.D. – 6120050961")
    idx = exay.sablonlari_indeksle(str(tmp_path))
    assert idx["6120050961"][0].endswith("oyak.docx")          # YMM yazısı da indekslenir


def test_birlesik_ymm_yazi_bloklara_ayrilir(tmp_path):
    """Tek dosyada birden çok YMM yazısı → 'Konu: Bilgi İsteme' başlığından bölünür."""
    import docx
    from copy import deepcopy
    from docx.oxml.ns import qn
    tek = tmp_path / "tek.docx"
    _ymm_yazi_docx_yaz(tek, "OYAK ÇİMENTO A.Ş.", "ANKARA KURUMLAR V.D. – 6120050961")
    d = docx.Document(str(tek))
    body = d.element.body
    els = [e for e in body if e.tag in (qn('w:p'), qn('w:tbl'))]
    sect = body.find(qn('w:sectPr'))
    for e in els:
        yeni = deepcopy(e)
        body.insert(list(body).index(sect), yeni) if sect is not None else body.append(yeni)
    bloklar = exay._docx_firma_bloklari(d)
    assert len(bloklar) == 2
    assert all(b["vkn"] == "6120050961" for b in bloklar)
